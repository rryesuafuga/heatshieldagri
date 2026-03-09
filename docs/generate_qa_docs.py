#!/usr/bin/env python3
"""
Generate updated HeatShield Agri QA documents (Professional and Talking Points).
Produces both .docx and .pdf versions of each document.

Key updates:
- Replace XGBoost with Random Forest throughout
- Update ML architecture description (physics-first, ONNX Runtime deployment)
- Update ML model status (now trained and deployed with 3 ONNX models)
- Add Raymond Reuel Wayesu / Wayesu Community Research Organisation Ltd branding
- Update environmental impact references
"""

import os
import copy
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from fpdf import FPDF

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))

AUTHOR = "Raymond Reuel Wayesu"
ORG = "Wayesu Community Research Organisation Ltd"
FOOTER_TEXT = f"{ORG} | {AUTHOR}"

# ============================================================================
# CONTENT DEFINITIONS
# ============================================================================

# ---------- QA Professional content ----------

QA_PRO_HEADER_INFO = (
    f"Prepared for: {AUTHOR}\n"
    f"Organisation: {ORG}\n"
    "Interview: March 9, 2026, 21:00 EAT\n"
    "Format: 5 min intro + 10 min pitch + 30 min Q&A"
)

QA_PRO_SECTIONS = [
    {
        "heading": "1. Technical & AI Questions",
        "qas": [
            {
                "q": "How accurate are your WBGT predictions?",
                "a": (
                    "Our predictions use the ISO 7243 WBGT calculation standard\u2014the internationally "
                    "recognized method for occupational heat assessment. We combine multiple data sources "
                    "including Open-Meteo weather forecasts (ECMWF and GFS models), SRTM terrain data, "
                    "and spatial interpolation algorithms. We display uncertainty bounds with every prediction "
                    "and use conservative thresholds that err toward caution. During the pilot phase, we'll "
                    "collect ground-truth data to quantify actual prediction error and continuously improve "
                    "calibration."
                ),
            },
            {
                "q": "Why did you choose Random Forest with a physics-first architecture?",
                "a": (
                    "Random Forest excels at tabular weather data\u2014temperature, humidity, wind speed, "
                    "terrain features\u2014capturing complex feature interactions with high interpretability "
                    "that agricultural stakeholders can trust. It is computationally efficient and requires "
                    "no GPU infrastructure, which is essential for sustainable operations in our context. "
                    "We deploy the trained models via ONNX Runtime, enabling cross-platform native C++ "
                    "inference on Android and browser-based inference on the web. The physics-first "
                    "architecture ensures strict ISO 7243 compliance: physics calculations provide the "
                    "foundational WBGT estimates, while Random Forest validates and enhances those "
                    "predictions using a 17-feature engineering pipeline (lag features, cyclical encodings, "
                    "rolling statistics, and delta calculations). This hybrid approach prevents physically "
                    "impossible predictions while leveraging ML to improve accuracy."
                ),
            },
            {
                "q": "Is the AI/ML model trained and deployed?",
                "a": (
                    "Yes\u2014the model is fully trained and deployed on both web and Android platforms. "
                    "We have 3 ONNX models in production: the core Random Forest WBGT prediction model, "
                    "a feature preprocessing model, and a post-processing calibration model. These are "
                    "served via ONNX Runtime for native C++ inference on Android and browser inference "
                    "on the web, ensuring fast, offline-capable predictions without requiring server-side "
                    "computation for each request."
                ),
            },
            {
                "q": "How does the work schedule optimizer function?",
                "a": (
                    "The optimizer analyzes the 72-hour WBGT forecast and identifies all hours below the "
                    "recommended work threshold for moderate agricultural work (ISO 7243 categories). It "
                    "applies practical constraints: farmers need 6-8 productive hours, prefer contiguous "
                    "blocks, and work during daylight. The output is specific recommendations like "
                    "'5:00 AM - 11:00 AM, then 4:00 PM - 6:30 PM' with estimated safe work hours per day. "
                    "This is a constraint satisfaction problem\u2014it's 'AI' in the sense of automated "
                    "intelligent analysis processing thousands of forecast combinations that humans couldn't "
                    "manually analyze."
                ),
            },
            {
                "q": "What happens when predictions are wrong?",
                "a": (
                    "We design for this scenario explicitly. Every alert includes guidance to 'listen to "
                    "your body' and recognize heat illness symptoms regardless of predictions. We use "
                    "conservative thresholds that minimize false negatives (predicted safe but actually "
                    "dangerous). We implement a feedback loop where users report via SMS whether conditions "
                    "matched predictions, enabling continuous model improvement. Systematic errors trigger "
                    "recalibration. The pilot phase explicitly validates prediction accuracy before scaling."
                ),
            },
        ],
    },
    {
        "heading": "2. Impact & Reach Questions",
        "qas": [
            {
                "q": "How will you reach farmers without smartphones?",
                "a": (
                    "This is exactly why USSD/SMS is a primary channel, not an afterthought. USSD works "
                    "on any basic feature phone without internet. Farmers dial *384*HEAT# and navigate "
                    "simple numbered menus to check heat risk, get forecasts, or register for alerts. SMS "
                    "messages are simple, actionable, in local languages: 'HIGH HEAT TODAY. Work 5AM-10AM "
                    "only. Rest midday. Drink water every 30 min.' For farmers without any phone, extension "
                    "workers and cooperative leaders share information through community gatherings. No "
                    "internet required\u2014USSD and SMS work on 2G networks covering nearly all of Uganda."
                ),
            },
            {
                "q": "How will you reach 10,000 farmers in 12 months?",
                "a": (
                    "We plug into existing infrastructure with trusted relationships. UNFFE's cooperative "
                    "networks in target districts become distribution partners\u2014we estimate 3-4 cooperatives "
                    "with 500-2000 members each. We train 10 extension workers who each work with ~100 "
                    "farmer groups, registering farmers during routine visits. Once registered, SMS delivery "
                    "is automated. Growth model: Months 1-3 = 500 farmers (proving the model), Months 4-6 "
                    "= 2,000 (pilot expansion), Months 7-12 = 10,000 (adding districts, word-of-mouth). "
                    "We grow through trusted networks, not cold outreach."
                ),
            },
            {
                "q": "How do you measure actual health impact?",
                "a": (
                    "Behavior change is our primary outcome\u2014measurable and causally linked to health "
                    "protection. We measure: (1) Self-reported work schedule changes, water intake, and "
                    "shade-seeking via monthly SMS check-ins and surveys; (2) Heat illness symptoms at "
                    "baseline and endline, comparing users vs. non-users; (3) Health facility data on "
                    "heat-related visits; (4) Productivity metrics including workdays lost. Our Makerere "
                    "University partnership ensures rigorous evaluation methodology with treatment and "
                    "comparison groups."
                ),
            },
            {
                "q": "How do you ensure equitable reach for women farmers?",
                "a": (
                    "Women face higher heat stress risk (PSI 3.28 vs 2.9 for men) but often have less "
                    "phone access. We partner with women's farming groups that have existing meeting "
                    "structures. For households where phones are shared or controlled by men, cooperative-level "
                    "alerts ensure women learn about heat risk through group meetings. We track registration "
                    "and engagement by gender and adjust outreach if gaps emerge. Our Farmer Advisory "
                    "Committee includes women farmers who guide message design and intervention strategies."
                ),
            },
        ],
    },
    {
        "heading": "3. Team & Execution Questions",
        "qas": [
            {
                "q": "Can you execute this as a solo founder?",
                "a": (
                    "The MVP was built solo, demonstrating execution capability. But the 12-month plan "
                    "doesn't rely on me doing everything alone. I handle technical development, ML models, "
                    "architecture, partnerships, and fundraising\u2014my core competencies from 10+ years of "
                    "research work. The $47K personnel budget funds a part-time developer and extension "
                    "worker payments. Partners contribute: UNFFE provides cooperative access, extension "
                    "workers do farmer training, Makerere provides research support. We're building a "
                    "system that scales through automation (SMS) and existing human infrastructure, not "
                    "through me personally reaching farmers."
                ),
            },
            {
                "q": "What qualifies you to lead this?",
                "a": (
                    f"{AUTHOR} brings a rare combination: 10+ years as biostatistician at Uganda Virus "
                    "Research Institute (health research and Uganda health systems expertise); MSc Statistics "
                    "from Linkoping University (rigorous statistical training); current PhD in Machine "
                    "Learning/Computer Vision (why I can build ML systems, not just talk about them); "
                    "technical stack proficiency in R, Python, Rust, TypeScript. I built the entire "
                    "HeatShield MVP myself. I understand health systems, statistical methods, ML "
                    "implementation, and Uganda's context\u2014a unique intersection that positions me for "
                    "this specific challenge."
                ),
            },
            {
                "q": "What's the status of your partnerships?",
                "a": (
                    "I want to be transparent: these are planned partnerships with preliminary conversations, "
                    "not signed MOUs. This is typical at pre-funding stage. UNFFE is interested but waiting "
                    "for funding confirmation. MAAIF (government) rarely commits before projects are funded. "
                    "Makerere has research collaboration interest pending project scope. Africa's Talking is "
                    "commercial\u2014pay-per-use, already integrated in MVP. Months 1-3 explicitly include "
                    "partnership formalization as a key deliverable. The relationships are real; the "
                    "paperwork comes after the grant."
                ),
            },
        ],
    },
    {
        "heading": "4. Sustainability & Scale Questions",
        "qas": [
            {
                "q": "What's your revenue model after the grant?",
                "a": (
                    "Basic alerts to individual farmers remain FREE forever\u2014this is a public good "
                    "subsidized by other revenue. B2B dashboards ($50-200/month) serve cooperatives, export "
                    "companies (EU supply chain compliance), and large farms needing documentation. "
                    "Institutional partnerships with Ministry of Agriculture, World Bank projects, and "
                    "USAID programs pay for tools their programs use. Climate adaptation funding from "
                    "Green Climate Fund, Adaptation Fund, and bilateral programs support exactly this "
                    "infrastructure. Cost efficiency advantage: $0.30/farmer/year for SMS means even "
                    "modest revenue covers massive reach."
                ),
            },
            {
                "q": "How does this scale beyond Uganda?",
                "a": (
                    "The architecture is designed for replication from day one. Weather data sources "
                    "(Open-Meteo, ICPAC) cover all East Africa. WBGT calculations are universal physics. "
                    "SMS/USSD infrastructure exists everywhere. Only configuration changes are needed for "
                    "new regions. East Africa has 100+ million agricultural workers facing similar "
                    "challenges\u2014Kenya, Tanzania, Rwanda, Ethiopia have the same climate patterns and "
                    "phone infrastructure. Our open-source strategy (MIT license) enables other organizations "
                    "to deploy in their contexts. We provide the toolkit; we don't gatekeep the technology."
                ),
            },
            {
                "q": "What if farmers don't change behavior despite alerts?",
                "a": (
                    "We don't assume automatic behavior change. Our approach: Economic framing shows "
                    "farmers how to maintain productivity while avoiding dangerous hours\u2014we say 'work "
                    "smarter' not 'don't work.' Trusted messengers (extension workers, cooperative leaders) "
                    "deliver information through pre-established relationships. Social proof builds as "
                    "early adopters report benefits. Iterative messaging adjusts based on what works\u2014if "
                    "text doesn't drive change, we add voice; if daily alerts cause fatigue, we adjust "
                    "frequency. Evidence from Kenya shows SMS agricultural advisories achieve 11.5% yield "
                    "increases\u2014behavior change through mobile messaging is proven."
                ),
            },
        ],
    },
    {
        "heading": "5. Responsible AI Questions",
        "qas": [
            {
                "q": "How do you handle farmer data privacy?",
                "a": (
                    "Data minimization is our core principle\u2014we collect only phone number (for SMS) "
                    "and district (for localized forecasts). No personal health information. Symptom surveys "
                    "are anonymized research data, not linked to phone numbers. Phone numbers are hashed "
                    "for storage, decrypted only for SMS sending. TLS encryption for all transmission. "
                    "Users actively opt-in with clear local-language explanation of data use. Easy opt-out "
                    "via SMS. User information is never sold or shared with third parties. We'll establish "
                    "a data governance committee including farmer representatives."
                ),
            },
            {
                "q": "How do you address AI bias?",
                "a": (
                    "We proactively address multiple bias types. Geographic bias: We use ICPAC regional "
                    "forecasts calibrated for East Africa and validate against local observations. Coverage "
                    "bias: Predictions cover all pilot districts equally, including remote sub-counties. "
                    "Access bias: USSD/SMS channels work on any phone, and cooperative-level sharing "
                    "reaches non-phone-owners. Gender bias: We track engagement by gender and adjust "
                    "outreach if women are underrepresented. Continuous monitoring: Bias testing is "
                    "ongoing\u2014we analyze whether prediction accuracy varies by district and adjust "
                    "accordingly."
                ),
            },
            {
                "q": "What about environmental impact of your AI?",
                "a": (
                    "Lightweight by design: Random Forest and spatial interpolation, not GPU-intensive "
                    "deep learning. The Random Forest models are deployed via ONNX Runtime for efficient "
                    "native inference\u2014no GPU required. Pre-calculated forecasts are cached and served, "
                    "avoiding on-demand inference for each request. The environmental cost is negligible "
                    "compared to climate adaptation benefits delivered\u2014we're helping millions adapt to "
                    "climate change. Cloud providers (Vercel, AWS) have renewable energy commitments. "
                    "We'll measure and report carbon footprint annually, and explore edge computing as "
                    "we scale."
                ),
            },
        ],
    },
    {
        "heading": "6. Strategic Questions",
        "qas": [
            {
                "q": "Why hasn't anyone built this already?",
                "a": (
                    "Several factors explain this gap: Climate services focus on rainfall (drought, flood)"
                    "\u2014heat stress falls between climate and health sectors with nobody's core mandate. "
                    "In smallholder agriculture, there's no employer to provide heat protection\u2014farmers "
                    "need institutional support. Previous approaches assumed smartphone apps, excluding most "
                    "farmers. This challenge requires climate science + ML + mobile tech + agricultural "
                    "extension + health research expertise\u2014a rare combination. VCs don't fund public "
                    "goods for poor farmers; development funders haven't prioritized occupational heat. "
                    "Programs like Activate AI fill exactly this gap."
                ),
            },
            {
                "q": "What if government wants to take over your system?",
                "a": (
                    "That would be success, not failure! Government adoption means sustainable, "
                    "institutionalized protection for all farmers. Our open-source code enables this\u2014if "
                    "MAAIF wants to integrate alerts into official services, they can. We'd become technical "
                    "advisors or pivot to other countries while Uganda runs its own system. Realistic "
                    "timeline: government tech adoption takes years, giving us time to demonstrate impact "
                    "and expand regionally. Our mission is protecting farmers, not building a company. "
                    "Government institutionalization achieves our mission at scale."
                ),
            },
            {
                "q": "What if you don't win this grant?",
                "a": (
                    "This grant is important but not existential. The MVP is built and running. I'm "
                    "pursuing IKI Small Grants (EUR 60-200K), ESTDEV (EUR 50-150K), UNICEF Venture Fund "
                    "($100K), and others\u2014multiple applications are active. Without major funding, I'd "
                    "run a smaller pilot with personal resources and UVRI networks. I could pursue "
                    "academic funding through my PhD program. Core work continues regardless: weather "
                    "data is free, hosting costs are minimal. Grant funding accelerates scale; it doesn't "
                    "determine existence."
                ),
            },
        ],
    },
]

QA_PRO_KEY_MESSAGES = [
    ("Working MVP", "Live at heatshieldagri.vercel.app -- proof of execution, not just plans"),
    ("Unserved Market", "12.4M agricultural workers with zero existing heat protection"),
    ("Proven Science", "ISO 7243 standard + Random Forest ML -- innovation is application"),
    ("Inclusive Design", "USSD/SMS for feature phones, local languages, gender equity focus"),
    ("Regional Scale", "Uganda proves concept for 100M+ East African workers"),
    ("Sustainability", "Freemium model: free farmer alerts, paid institutional dashboards"),
    ("Open Source", "MIT-licensed algorithms enable replication without us"),
    ("Founder Fit", "Rare intersection: biostatistics + ML + Uganda health systems expertise"),
]

# ---------- QA Talking Points content ----------

TP_SECTIONS = [
    {
        "heading": "SECTION 1: TECHNICAL & AI QUESTIONS",
        "qas": [
            {
                "q": "How accurate are your WBGT predictions? What's your error rate?",
                "points": [
                    "We use the ISO 7243 WBGT calculation standard, which is the internationally recognized method for occupational heat assessment. This isn't a custom algorithm\u2014it's proven science.",
                    "Our predictions combine multiple data sources: Open-Meteo weather forecasts (which draw from ECMWF and GFS global models), terrain data from SRTM, and spatial interpolation using Inverse Distance Weighting.",
                    "Current MVP validation: We cross-reference our WBGT calculations against expected values from known weather conditions. The physics-based approach means our error comes primarily from upstream weather forecast uncertainty, not our calculations.",
                    "Planned improvements: Integration with ICPAC regional forecasts calibrated for East Africa will improve accuracy. We'll also collect ground-truth data during pilot to quantify actual prediction error.",
                    "Importantly, we display uncertainty bounds with every prediction\u2014we never present point estimates as certainties. We also use conservative thresholds that err toward caution.",
                ],
                "key_msg": "Our predictions are based on international standards and multiple data sources. We're transparent about uncertainty and plan rigorous validation during the pilot.",
            },
            {
                "q": "Why did you choose Random Forest with a physics-first architecture?",
                "points": [
                    "Random Forest excels at tabular weather data with complex feature interactions (temperature, humidity, wind speed, terrain). It offers high interpretability that agricultural stakeholders can trust\u2014each prediction can be traced back to feature importance scores.",
                    "The physics-first architecture ensures ISO 7243 compliance: physics calculations provide the foundational WBGT estimates, while Random Forest validates and enhances those predictions. This prevents the model from making physically impossible predictions\u2014a common failure mode of pure ML approaches.",
                    "We use a 17-feature engineering pipeline: lag features capture temporal patterns, cyclical encodings handle time-of-day and seasonal effects, rolling statistics smooth noise, and delta calculations detect rapid weather changes.",
                    "Random Forest is computationally efficient and deployable without GPU infrastructure. We deploy via ONNX Runtime for native C++ inference on Android and browser-based inference on the web\u2014enabling fast, offline-capable predictions.",
                    "This architecture is well-established in climate science applications\u2014we're applying proven methods to an underserved use case, not inventing new ML techniques.",
                ],
                "key_msg": "We chose Random Forest for its interpretability, efficiency, and proven performance on tabular data. ONNX Runtime deployment enables cross-platform native inference without GPU.",
            },
            {
                "q": "Is the AI/ML model trained and deployed?",
                "points": [
                    "Yes\u2014the model is fully trained and deployed on both web and Android platforms.",
                    "We have 3 ONNX models in production: the core Random Forest WBGT prediction model, a feature preprocessing model, and a post-processing calibration model.",
                    "ONNX Runtime enables native C++ inference on Android and browser-based inference on the web, ensuring fast predictions without server-side computation for each request.",
                    "The 17-feature engineering pipeline processes raw weather data through lag features, cyclical encodings, rolling statistics, and delta calculations before feeding into the Random Forest model.",
                    "Offline capability means farmers in areas with intermittent connectivity can still receive predictions once weather data is cached locally.",
                ],
                "key_msg": "The ML model is fully trained and deployed with 3 ONNX models running on both web and Android. This is production-ready, not a prototype.",
            },
            {
                "q": "How does your AI work schedule optimizer function?",
                "points": [
                    "The optimizer analyzes the 72-hour WBGT forecast and identifies all hours below the recommended work threshold for the specific work intensity (we use ISO 7243 categories for moderate agricultural work).",
                    "It then applies constraints: farmers typically need 6-8 hours of productive work, prefer contiguous blocks, and work patterns should align with daylight hours.",
                    "The algorithm outputs recommended work windows (e.g., '5:00 AM - 11:00 AM, then 4:00 PM - 6:30 PM') with estimated safe work hours per day.",
                    "This is a constraint satisfaction problem, not complex ML\u2014but presenting it as 'AI-optimized' is accurate because the recommendations emerge from processing thousands of forecast data points that humans couldn't manually analyze.",
                ],
                "key_msg": "The optimizer is practical and explainable\u2014it finds safe hours within farmer constraints. It's AI in the sense of automated intelligent analysis, not black-box magic.",
            },
            {
                "q": "What happens when your system makes a wrong prediction?",
                "points": [
                    "First, we design for this scenario explicitly. Every alert includes guidance to 'listen to your body' and recognize heat illness symptoms regardless of what the system says.",
                    "False negative (predicted safe but actually dangerous): Our conservative thresholds minimize this risk. We'd rather over-warn than under-warn. If a false negative is reported, we adjust thresholds upward for that region.",
                    "False positive (predicted dangerous but actually safe): This is less harmful\u2014some lost productivity\u2014but still problematic for trust. We track user feedback on whether conditions matched predictions and use this to improve calibration.",
                    "We implement a feedback loop: Users can report whether alerts matched reality via SMS. Systematic errors trigger model recalibration.",
                    "The pilot phase is explicitly designed to validate prediction accuracy against real conditions before scaling.",
                ],
                "key_msg": "We design for failure, use conservative thresholds, collect feedback, and continuously improve. No prediction system is perfect\u2014the question is how you handle errors.",
            },
        ],
    },
    {
        "heading": "SECTION 2: IMPACT & REACH QUESTIONS",
        "qas": [
            {
                "q": "How will you reach farmers who don't have smartphones?",
                "points": [
                    "This is exactly why we built USSD/SMS as a primary channel, not an afterthought. USSD works on any basic phone\u2014the $15 feature phones that most Ugandan farmers have.",
                    "USSD menu: Farmers dial *384*HEAT# (or similar short code) and navigate simple numbered menus to check today's heat risk, get 3-day forecasts, or register for SMS alerts.",
                    "SMS alerts: Registered farmers receive automated alerts when WBGT exceeds their threshold. Messages are simple, actionable, in local languages: 'HIGH HEAT TODAY. Work 5AM-10AM only. Rest midday. Drink water every 30 min.'",
                    "Community amplification: For farmers without any phone, extension workers and cooperative leaders receive alerts and share information through existing community gatherings. Heat warnings become part of cooperative announcements.",
                    "No internet required: USSD and SMS work on 2G networks, which cover nearly all of Uganda including remote areas.",
                ],
                "key_msg": "We meet farmers where they are\u2014feature phones, no internet, local languages. Technology should adapt to users, not vice versa.",
            },
            {
                "q": "Your target is 10,000 farmers in 12 months. How will you achieve this?",
                "points": [
                    "We're not building a user acquisition machine from scratch\u2014we're plugging into existing infrastructure with trusted relationships.",
                    "Cooperative partnership: UNFFE (Uganda National Farmers Federation) has established cooperative networks in our target districts. They become distribution partners, not just endorsers. We estimate 3-4 cooperatives with 500-2000 members each.",
                    "Extension worker network: We train 10 extension workers who each work with ~100 farmer groups. They register farmers during routine visits and train group leaders to cascade the information.",
                    "SMS broadcast efficiency: Once a farmer is registered (name, phone, district), delivery is automated. We're not doing 10,000 individual outreach\u2014we're enabling organic spread through trusted channels.",
                    "Growth model: Months 1-3 = 500 farmers (proving model works), Months 4-6 = 2,000 (expansion within pilot districts), Months 7-12 = 10,000 (adding districts, word-of-mouth).",
                ],
                "key_msg": "We grow through existing trusted networks, not cold outreach. 10,000 is ambitious but achievable through partnerships already in discussion.",
            },
            {
                "q": "How do you measure actual health impact, not just engagement?",
                "points": [
                    "Behavior change is our primary outcome\u2014it's measurable and causally linked to health protection even without directly measuring illness rates.",
                    "Behavior metrics: Self-reported work schedule changes, water intake changes, shade-seeking behavior. Measured via monthly SMS check-ins and quarterly surveys.",
                    "Health symptom tracking: Baseline and endline surveys measuring self-reported heat illness symptoms (headache, dizziness, nausea, fatigue during hot periods). We compare users vs. non-users in same communities.",
                    "Health facility data: Partnership with local health centers to track heat-related visits. This is supplementary\u2014we don't expect large numbers given our intervention is preventive.",
                    "Productivity metrics: Farmer-reported workdays lost to heat illness; cooperative-level productivity data where available.",
                    "Research partnership with Makerere University ensures methodological rigor\u2014they'll help design the evaluation framework.",
                ],
                "key_msg": "We measure the full chain: alert delivery -> behavior change -> symptom reduction. Makerere partnership ensures rigorous evaluation design.",
            },
            {
                "q": "What about women farmers? How do you ensure equitable reach?",
                "points": [
                    "We know women face HIGHER heat stress risk (research shows PSI 3.28 vs 2.9 for men) but often have LESS access to phones and information. This is a design priority, not an afterthought.",
                    "Women-specific outreach: We partner with women's farming groups and women-focused cooperatives. These groups have existing meeting structures where information can be shared.",
                    "Phone access reality: In many households, phones are shared or controlled by men. Our cooperative-level alerts ensure women learn about heat risk through group meetings even if they don't personally receive SMS.",
                    "Gender-disaggregated monitoring: We track registration and engagement rates by gender. If we see gaps, we adjust outreach strategies.",
                    "Farmer Advisory Committee: Our 10-member advisory group explicitly includes women farmers who guide how we design messages and interventions.",
                ],
                "key_msg": "Women's higher vulnerability requires intentional design. We partner with women's groups and track gender equity metrics throughout.",
            },
        ],
    },
    {
        "heading": "SECTION 3: TEAM & EXECUTION QUESTIONS",
        "qas": [
            {
                "q": "Can you execute this as a solo founder?",
                "points": [
                    "I want to address this directly: Yes, the MVP was built solo, which actually demonstrates execution capability. But the 12-month plan doesn't rely on me doing everything alone.",
                    "What I do: Technical development, ML model training, system architecture, strategic partnerships, fundraising. These are my core competencies from 10+ years of research work.",
                    "What the budget funds: The $47K personnel budget includes hiring a part-time developer for Android app and support from existing extension workers who are paid per activity.",
                    "Partner contributions: UNFFE provides cooperative access and mobilization. Extension workers (MAAIF network) do farmer training. Makerere provides research support. I orchestrate, I don't do it all.",
                    "Scaling approach: We're building a system that scales through automation (SMS alerts) and existing human infrastructure (extension workers), not through me personally reaching farmers.",
                ],
                "key_msg": "Solo founder does not equal solo execution. I built the MVP; the grant funds a team and leverages partner infrastructure for field deployment.",
            },
            {
                "q": "What's your background? Why are you qualified to lead this?",
                "points": [
                    f"10+ years as biostatistician at Uganda Virus Research Institute (UVRI). {AUTHOR} understands health research, epidemiology, and working within Uganda's health system.",
                    "MSc Statistics from Linkoping University (Sweden). Rigorous training in statistical modeling, uncertainty quantification, and data analysis.",
                    "Current PhD in Machine Learning/Computer Vision. This is why I can build ML systems, not just talk about them.",
                    "Technical stack experience: R, Python, Rust, TypeScript. I built the entire HeatShield MVP myself\u2014proof of technical capability.",
                    f"Domain expertise intersection: {AUTHOR} understands health systems (from UVRI), statistical methods (from training), ML implementation (from PhD), and Uganda's context (from living and working here). This combination is rare.",
                ],
                "key_msg": "I combine statistical rigor, ML capability, health research experience, and deep Uganda context. The working MVP demonstrates I can execute, not just plan.",
            },
            {
                "q": "What's the status of your partnerships? Are they confirmed?",
                "points": [
                    "I want to be transparent: These are planned partnerships with preliminary conversations, not signed MOUs. This is typical at pre-funding stage.",
                    "UNFFE: Initial discussions about cooperative access for pilot districts. They're interested but waiting to see funding confirmation before formal commitment.",
                    "MAAIF: Exploratory conversations about extension worker involvement. Government agencies rarely commit before projects are funded.",
                    "Makerere University School of Public Health: Research collaboration discussions. They have interest in climate-health research but need to see project scope.",
                    "Africa's Talking: This is purely commercial\u2014they provide SMS/USSD infrastructure on a pay-per-use basis. No partnership needed, just technical integration (already working in MVP).",
                    "Months 1-3 of the project explicitly include partnership formalization as a key deliverable.",
                ],
                "key_msg": "Partnerships are in discussion, awaiting funding to formalize. This is normal pre-award stage. The relationships are real; the paperwork comes after the grant.",
            },
        ],
    },
    {
        "heading": "SECTION 4: SUSTAINABILITY & SCALE QUESTIONS",
        "qas": [
            {
                "q": "What's your revenue model? How do you sustain this after the grant?",
                "points": [
                    "Basic alerts to individual farmers remain FREE forever. This is a public good, and we subsidize it through other revenue streams.",
                    "B2B dashboards ($50-200/month): Agricultural cooperatives, export companies (EU supply chain compliance), and large farms want dashboards showing worker heat exposure. They'll pay for documentation and management tools.",
                    "Institutional partnerships: Ministry of Agriculture may integrate heat alerts into official extension services. Development partners (World Bank projects, USAID programs) often pay for tools their programs use.",
                    "Climate adaptation funding: Green Climate Fund, Adaptation Fund, and bilateral climate finance programs fund exactly this kind of climate resilience infrastructure. We'll pursue these grants.",
                    "Cost efficiency advantage: Our core cost is $0.30/farmer/year for SMS. Even modest institutional support or B2B revenue covers massive farmer reach.",
                ],
                "key_msg": "Freemium model: free alerts for farmers, paid dashboards for institutions. Low costs mean even small revenue streams enable sustainability.",
            },
            {
                "q": "How does this scale beyond Uganda?",
                "points": [
                    "The architecture is designed for replication from day one. We're not building Uganda-specific systems that would require rebuilding for each country.",
                    "Technical scalability: Weather data sources (Open-Meteo, ICPAC) cover all of East Africa. WBGT calculations are universal physics. SMS/USSD work everywhere. Only configuration changes needed for new regions.",
                    "Regional scope: East Africa has 100+ million agricultural workers facing similar heat challenges. Kenya, Tanzania, Rwanda, Ethiopia\u2014same climate patterns, same phone infrastructure, same need.",
                    "Open source strategy: Core algorithms and integration templates on GitHub (MIT license). Other organizations can deploy in their contexts. We provide the toolkit, not gatekeep the technology.",
                    "Pan-African expansion path: Success in Uganda -> ICPAC partnership for regional coverage -> implementation partners in neighboring countries. We don't have to do it all ourselves.",
                ],
                "key_msg": "Uganda is proof of concept for a regional solution. Open source + ICPAC partnership enables scale across 100M+ East African agricultural workers.",
            },
            {
                "q": "What if farmers don't change their behavior despite receiving alerts?",
                "points": [
                    "This is a real risk, and we don't assume behavior change happens automatically. Our approach addresses it directly.",
                    "Economic framing: Work schedule optimization shows farmers how to maintain productivity while avoiding dangerous hours. We're not saying 'don't work'\u2014we're saying 'work smarter.'",
                    "Trusted messengers: Information comes through extension workers and cooperative leaders, not an unknown app. Trust is pre-established.",
                    "Social proof: As early adopters report benefits (less fatigue, fewer symptoms), word-of-mouth builds adoption. Farmer-to-farmer influence is powerful.",
                    "Iterative messaging: If initial messages don't drive behavior change, we adjust. Maybe text isn't working\u2014we add voice. Maybe daily alerts cause fatigue\u2014we adjust frequency. The pilot is for learning, not just deploying.",
                    "Research evidence: Similar interventions (SMS agricultural advisories) show 11.5% yield increases in Kenya. Behavior change through mobile messaging is proven; we adapt it for heat.",
                ],
                "key_msg": "We design for behavior change: economic framing, trusted messengers, iterative improvement. The pilot will test and refine our approach.",
            },
        ],
    },
    {
        "heading": "SECTION 5: RESPONSIBLE AI QUESTIONS",
        "qas": [
            {
                "q": "How do you handle data privacy for farmers?",
                "points": [
                    "Data minimization is our core principle. We collect only: phone number (for SMS delivery) and district (for localized forecasts). That's it.",
                    "No health data: We don't collect personal health information. Symptom surveys are anonymized research data, not linked to individual phone numbers.",
                    "Encryption: Phone numbers are hashed for storage and only decrypted for SMS sending. TLS encryption for all data transmission.",
                    "Consent: Users actively opt-in by registering. Clear explanation in local languages of what we collect and why. Easy opt-out via SMS.",
                    "No data sales: User information is never sold or shared with third parties. Period. This isn't a data harvesting operation.",
                    "Local governance: We'll establish a data governance committee including farmer representatives to oversee data practices.",
                ],
                "key_msg": "We collect minimal data, encrypt it, get clear consent, and never sell it. Farmers' privacy is protected by design.",
            },
            {
                "q": "What about AI bias? Could your system disadvantage certain farmers?",
                "points": [
                    "Weather model bias: Global weather models may be less accurate in data-sparse African regions. We mitigate by using ICPAC regional forecasts calibrated for East Africa and validating against local observations.",
                    "Geographic coverage bias: We ensure predictions cover all pilot districts equally, including remote sub-counties. We don't just serve easy-to-reach areas.",
                    "Access bias: If only tech-savvy, phone-owning farmers benefit, we'd widen inequalities. That's why USSD/SMS channels (works on any phone) and cooperative-level information sharing are core to our design.",
                    "Gender bias: We track registration and engagement by gender. If women are underrepresented, we adjust outreach strategies to ensure equitable access.",
                    "Continuous monitoring: Bias testing is ongoing, not one-time. We'll analyze whether prediction accuracy varies by district and adjust accordingly.",
                ],
                "key_msg": "We proactively address geographic, access, and gender bias through inclusive design and continuous monitoring. Equity is a design requirement.",
            },
            {
                "q": "What's the environmental impact of your AI system?",
                "points": [
                    "Lightweight by design: We use Random Forest and spatial interpolation, not large language models or GPU-intensive deep learning. Standard cloud infrastructure, minimal energy footprint.",
                    "Efficient architecture: The Random Forest models are deployed via ONNX Runtime for native inference\u2014no GPU required. Pre-calculated forecasts are cached and served to users. We don't run inference on-demand for each request\u2014that would be computationally wasteful.",
                    "Proportionate impact: The environmental cost of our AI system is negligible compared to the climate adaptation benefits delivered. We're helping millions adapt to climate change; our compute costs are a rounding error.",
                    "Cloud providers: Vercel and AWS (our infrastructure) have committed to renewable energy. AWS Africa region uses increasingly renewable power.",
                    "Future commitment: We'll measure and report our carbon footprint annually. As we scale, we'll explore edge computing to reduce data center dependency.",
                ],
                "key_msg": "Our AI is lightweight and efficient. The adaptation benefits far outweigh minimal compute costs. We'll track and report our footprint.",
            },
        ],
    },
    {
        "heading": "SECTION 6: POTENTIAL CURVEBALL QUESTIONS",
        "qas": [
            {
                "q": "Why hasn't anyone else built this already?",
                "points": [
                    "Great question. There are a few reasons this gap exists:",
                    "Climate services focus: Most agricultural climate services focus on rainfall (drought, flood prediction). Heat stress is an occupational health issue that falls between climate and health sectors\u2014nobody's core mandate.",
                    "Occupational health gap: In high-income countries, employers provide heat protection. In smallholder agriculture, there's no employer\u2014farmers are their own bosses with no institutional protection.",
                    "Technology assumption: Previous approaches assumed smartphone apps, which excludes most farmers. USSD/SMS delivery for heat warnings is technically straightforward but requires someone to build it.",
                    "Interdisciplinary challenge: This needs climate science + ML + mobile technology + agricultural extension + health research expertise. That combination is rare. My background uniquely positions me here.",
                    "Market viability: VCs don't fund public goods for poor farmers. Development funders haven't prioritized occupational heat. Programs like Activate AI are exactly what's needed to fill this gap.",
                ],
                "key_msg": "It's a coordination failure across sectors, not a technical impossibility. The technology exists; it just hasn't been applied to this population.",
            },
            {
                "q": "What if government policy changes and they want to take over your system?",
                "points": [
                    "That would be a success, not a failure! Government adoption means sustainable, institutionalized heat protection for all farmers.",
                    "Open source design: Our code is openly available. If MAAIF wants to integrate heat alerts into official extension services, they can. We'd help them.",
                    "Transition path: We could become technical advisors, training government staff to maintain the system. Or we pivot to other countries while Uganda runs its own system.",
                    "Realistic timeline: Government tech adoption takes years. We'd have plenty of time to demonstrate impact, build our reputation, and expand regionally before any takeover scenario.",
                    "Mission alignment: Our goal is protecting farmers, not building a company. Government institutionalization achieves our mission at scale.",
                ],
                "key_msg": "Government adoption is the best-case scenario for impact. Our open-source approach facilitates rather than prevents this.",
            },
            {
                "q": "What would you do if you don't win this grant?",
                "points": [
                    "This grant is important but not existential. The MVP is built and running. I'll continue regardless.",
                    "Alternative funding: I'm pursuing IKI Small Grants (EUR 60-200K, deadline passed but future rounds), ESTDEV (EUR 50-150K), UNICEF Venture Fund ($100K), and others. This isn't my only application.",
                    "Smaller scale start: Without major funding, I'd run a smaller pilot with personal resources and existing UVRI networks. Slower but still progress.",
                    "Research angle: I could pursue academic funding through my PhD program or UVRI collaborations. A research-framed pilot might access different funding streams.",
                    "The core work continues: Weather data is free. Hosting costs are minimal. I can maintain and improve the platform regardless. Grant funding accelerates scale, but doesn't determine existence.",
                ],
                "key_msg": "I'm pursuing multiple funding paths. HeatShield Agri continues regardless\u2014this grant accelerates scale but isn't a binary success/failure point.",
            },
        ],
    },
]

TP_FINAL_NOTES = {
    "themes": [
        "Working MVP -- You've built it, it's live, reviewers can see it. This is rare and valuable.",
        "Unserved population -- 12.4M agricultural workers with zero existing heat protection. Clear market gap.",
        "Proven science -- ISO 7243 standard, Random Forest ML methods. Innovation is application, not algorithms.",
        "Inclusive design -- USSD/SMS for feature phones, local languages, gender equity. Tech that meets users where they are.",
        "Regional scalability -- Uganda is proof of concept for 100M+ East African agricultural workers.",
    ],
    "avoid": [
        "Don't oversell ML sophistication -- be honest about using proven methods appropriately",
        "Don't claim confirmed partnerships when they're in discussion -- transparency builds trust",
        "Don't be defensive about solo founder status -- own it as proof of execution capability",
        "Don't forget to demo the MVP -- always offer to show the live platform",
    ],
    "checklist": [
        "Test internet connection and backup (mobile hotspot)",
        "Have heatshieldagri.vercel.app open and ready to share screen",
        "Quiet environment, professional background",
        "Water nearby (long interview!)",
        "This document for reference (but don't read from it)",
        "Join Zoom link 5 minutes early",
    ],
}


# ============================================================================
# DOCX GENERATION
# ============================================================================


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    shading_elm.append(shading)


def add_footer(doc, text):
    """Add a footer to the document."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)


def generate_qa_professional_docx(output_path):
    """Generate the QA Professional .docx file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HEATSHIELD AGRI")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 100, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Powered Heat Stress Early Warning System")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTERVIEW Q&A REFERENCE")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Activate AI: Economic Opportunity Challenge")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    # Header info table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = QA_PRO_HEADER_INFO
    set_cell_shading(cell, "F0F0F0")

    doc.add_paragraph()

    # Q&A sections
    for section in QA_PRO_SECTIONS:
        doc.add_heading(section["heading"], level=1)

        for qa in section["qas"]:
            # Question-Answer table
            qa_table = doc.add_table(rows=2, cols=1)
            qa_table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Question row
            q_cell = qa_table.cell(0, 0)
            q_cell.text = ""
            q_para = q_cell.paragraphs[0]
            q_run = q_para.add_run(qa["q"])
            q_run.bold = True
            q_run.font.size = Pt(11)
            set_cell_shading(q_cell, "E8F5E9")

            # Answer row
            a_cell = qa_table.cell(1, 0)
            a_cell.text = qa["a"]
            for para in a_cell.paragraphs:
                para.style.font.size = Pt(10)

            doc.add_paragraph()

    # Key Messages Summary
    doc.add_heading("Key Messages Summary", level=1)

    km_table = doc.add_table(rows=len(QA_PRO_KEY_MESSAGES) + 1, cols=2)
    km_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_theme = km_table.cell(0, 0)
    hdr_theme.text = ""
    p = hdr_theme.paragraphs[0]
    run = p.add_run("Theme")
    run.bold = True
    set_cell_shading(hdr_theme, "2E7D32")
    run.font.color.rgb = RGBColor(255, 255, 255)

    hdr_msg = km_table.cell(0, 1)
    hdr_msg.text = ""
    p = hdr_msg.paragraphs[0]
    run = p.add_run("Key Message")
    run.bold = True
    set_cell_shading(hdr_msg, "2E7D32")
    run.font.color.rgb = RGBColor(255, 255, 255)

    for i, (theme, msg) in enumerate(QA_PRO_KEY_MESSAGES):
        km_table.cell(i + 1, 0).text = theme
        km_table.cell(i + 1, 1).text = msg
        # Alternate shading
        if i % 2 == 0:
            set_cell_shading(km_table.cell(i + 1, 0), "F5F5F5")
            set_cell_shading(km_table.cell(i + 1, 1), "F5F5F5")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Remember: You've built something real. Show them.")
    run.bold = True
    run.font.size = Pt(12)

    # Footer
    add_footer(doc, FOOTER_TEXT)

    doc.save(output_path)
    print(f"  Created: {output_path}")


def generate_qa_talking_points_docx(output_path):
    """Generate the QA Talking Points .docx file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ACTIVATE AI CHALLENGE")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 100, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Finalist Interview: Q&A Preparation")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Detailed Talking Points for {AUTHOR}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ORG)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.italic = True

    # Interview structure
    p = doc.add_paragraph()
    run = p.add_run("INTERVIEW STRUCTURE")
    run.bold = True
    run.font.size = Pt(13)

    for item in [
        ("5 minutes:", "Introductions"),
        ("10 minutes:", "Pitch the proposal"),
        ("30 minutes:", "Q&A session"),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"  * {item[0]} ")
        run.bold = True
        p.add_run(item[1])

    # Sections
    for section in TP_SECTIONS:
        doc.add_heading(section["heading"], level=1)

        for qa in section["qas"]:
            # Question
            p = doc.add_paragraph()
            run = p.add_run(f"Q: {qa['q']}")
            run.bold = True
            run.font.size = Pt(11)

            # Talking Points header
            p = doc.add_paragraph()
            run = p.add_run("Talking Points:")
            run.bold = True

            # Numbered points
            for idx, point in enumerate(qa["points"], 1):
                p = doc.add_paragraph()
                p.add_run(f"{idx}. {point}")

            # Key Message box
            p = doc.add_paragraph()
            run = p.add_run(f"Key Message: {qa['key_msg']}")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)

            doc.add_paragraph()  # spacing

    # Final Preparation Notes
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("FINAL PREPARATION NOTES")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("Key Themes to Emphasize Throughout:")
    run.bold = True

    for idx, theme in enumerate(TP_FINAL_NOTES["themes"], 1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. {theme}")

    p = doc.add_paragraph()
    run = p.add_run("Things to Avoid:")
    run.bold = True

    for item in TP_FINAL_NOTES["avoid"]:
        p = doc.add_paragraph()
        p.add_run(f"  * {item}")

    p = doc.add_paragraph()
    run = p.add_run("Pre-Interview Checklist:")
    run.bold = True

    for item in TP_FINAL_NOTES["checklist"]:
        p = doc.add_paragraph()
        p.add_run(f"  [ ] {item}")

    p = doc.add_paragraph()
    run = p.add_run(f"Good luck, Raymond! You've done the work. Now show them what you've built.")
    run.bold = True
    run.font.size = Pt(12)

    # Footer
    add_footer(doc, FOOTER_TEXT)

    doc.save(output_path)
    print(f"  Created: {output_path}")


# ============================================================================
# PDF GENERATION
# ============================================================================


class QAPdf(FPDF):
    """Custom FPDF class with header/footer for HeatShield QA docs."""

    def __init__(self, doc_title=""):
        super().__init__()
        self.doc_title = doc_title
        # Register DejaVu for Unicode support
        self.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
        self.add_font("DejaVu", "B", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")

    def header(self):
        self.set_font("DejaVu", "B", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 5, "HeatShield Agri", 0, 0, "L")
        self.cell(0, 5, self.doc_title, 0, 1, "R")
        self.set_draw_color(200, 200, 200)
        self.line(10, 12, self.w - 10, 12)
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font("DejaVu", "", 7)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, FOOTER_TEXT, 0, 0, "L")
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", 0, 0, "R")

    def section_heading(self, text):
        self.set_font("DejaVu", "B", 14)
        self.set_text_color(46, 125, 50)  # dark green
        self.ln(4)
        self.cell(0, 8, text, 0, 1)
        self.set_draw_color(46, 125, 50)
        self.line(10, self.get_y(), self.w - 10, self.get_y())
        self.ln(4)

    def question_box(self, text):
        self.set_fill_color(232, 245, 233)  # light green
        self.set_font("DejaVu", "B", 11)
        self.set_text_color(0, 0, 0)
        x = self.get_x()
        w = self.w - 20
        self.multi_cell(w, 6, text, 0, "L", True)
        self.ln(1)

    def answer_text(self, text):
        self.set_font("DejaVu", "", 10)
        self.set_text_color(40, 40, 40)
        self.multi_cell(self.w - 20, 5.5, text, 0, "L")
        self.ln(3)

    def key_message_box(self, text):
        self.set_fill_color(255, 248, 225)  # light yellow
        self.set_font("DejaVu", "B", 9)
        self.set_text_color(46, 125, 50)
        self.multi_cell(self.w - 20, 5.5, f"Key Message: {text}", 0, "L", True)
        self.ln(4)

    def numbered_point(self, num, text):
        self.set_font("DejaVu", "", 10)
        self.set_text_color(40, 40, 40)
        self.set_x(self.l_margin)
        self.multi_cell(self.w - self.l_margin - self.r_margin, 5.5, f"{num}. {text}", 0, "L")
        self.ln(1)


def generate_qa_professional_pdf(output_path):
    """Generate the QA Professional .pdf file."""
    pdf = QAPdf("Interview Q&A Reference")
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Title
    pdf.set_font("DejaVu", "B", 22)
    pdf.set_text_color(0, 100, 0)
    pdf.cell(0, 12, "HEATSHIELD AGRI", 0, 1, "C")

    pdf.set_font("DejaVu", "", 13)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 8, "AI-Powered Heat Stress Early Warning System", 0, 1, "C")

    pdf.set_font("DejaVu", "B", 16)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 10, "INTERVIEW Q&A REFERENCE", 0, 1, "C")

    pdf.set_font("DejaVu", "", 11)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 7, "Activate AI: Economic Opportunity Challenge", 0, 1, "C")

    pdf.ln(3)

    # Info box
    pdf.set_fill_color(240, 240, 240)
    pdf.set_font("DejaVu", "", 10)
    pdf.set_text_color(40, 40, 40)
    pdf.multi_cell(0, 5.5, QA_PRO_HEADER_INFO, 0, "L", True)
    pdf.ln(5)

    # Sections
    for section in QA_PRO_SECTIONS:
        pdf.section_heading(section["heading"])
        for qa in section["qas"]:
            pdf.question_box(qa["q"])
            pdf.answer_text(qa["a"])

    # Key Messages Summary
    pdf.section_heading("Key Messages Summary")

    # Table header
    pdf.set_fill_color(46, 125, 50)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("DejaVu", "B", 10)
    pdf.cell(45, 7, "Theme", 1, 0, "L", True)
    pdf.cell(0, 7, "Key Message", 1, 1, "L", True)

    # Table rows
    pdf.set_text_color(0, 0, 0)
    for i, (theme, msg) in enumerate(QA_PRO_KEY_MESSAGES):
        if i % 2 == 0:
            pdf.set_fill_color(245, 245, 245)
        else:
            pdf.set_fill_color(255, 255, 255)
        pdf.set_font("DejaVu", "B", 9)
        pdf.cell(45, 6.5, theme, 1, 0, "L", True)
        pdf.set_font("DejaVu", "", 9)
        pdf.cell(0, 6.5, msg, 1, 1, "L", True)

    pdf.ln(6)
    pdf.set_font("DejaVu", "B", 12)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, "Remember: You've built something real. Show them.", 0, 1, "C")

    pdf.output(output_path)
    print(f"  Created: {output_path}")


def generate_qa_talking_points_pdf(output_path):
    """Generate the QA Talking Points .pdf file."""
    pdf = QAPdf("Interview Talking Points")
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Title
    pdf.set_font("DejaVu", "B", 20)
    pdf.set_text_color(0, 100, 0)
    pdf.cell(0, 12, "ACTIVATE AI CHALLENGE", 0, 1, "C")

    pdf.set_font("DejaVu", "", 13)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, "Finalist Interview: Q&A Preparation", 0, 1, "C")

    pdf.set_font("DejaVu", "", 11)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 7, f"Detailed Talking Points for {AUTHOR}", 0, 1, "C")

    pdf.set_font("DejaVu", "", 10)
    pdf.cell(0, 6, ORG, 0, 1, "C")

    pdf.ln(4)

    # Interview structure
    pdf.set_font("DejaVu", "B", 13)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, "INTERVIEW STRUCTURE", 0, 1, "L")

    pdf.set_font("DejaVu", "", 10)
    pdf.set_text_color(40, 40, 40)
    for item in [
        ("5 minutes:", "Introductions"),
        ("10 minutes:", "Pitch the proposal"),
        ("30 minutes:", "Q&A session"),
    ]:
        pdf.set_font("DejaVu", "B", 10)
        pdf.cell(5, 6, "", 0, 0)
        pdf.cell(25, 6, item[0], 0, 0)
        pdf.set_font("DejaVu", "", 10)
        pdf.cell(0, 6, item[1], 0, 1)

    pdf.ln(4)

    # Sections
    for section in TP_SECTIONS:
        pdf.section_heading(section["heading"])

        for qa in section["qas"]:
            pdf.question_box(f"Q: {qa['q']}")

            # Talking Points header
            pdf.set_font("DejaVu", "B", 10)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(0, 6, "Talking Points:", 0, 1)

            for idx, point in enumerate(qa["points"], 1):
                pdf.numbered_point(idx, point)

            pdf.key_message_box(qa["key_msg"])

    # Final Preparation Notes
    pdf.add_page()
    pdf.set_font("DejaVu", "B", 14)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 10, "FINAL PREPARATION NOTES", 0, 1, "L")
    pdf.ln(2)

    pdf.set_font("DejaVu", "B", 11)
    pdf.cell(0, 7, "Key Themes to Emphasize Throughout:", 0, 1)

    for idx, theme in enumerate(TP_FINAL_NOTES["themes"], 1):
        pdf.numbered_point(idx, theme)

    pdf.ln(3)
    pdf.set_font("DejaVu", "B", 11)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 7, "Things to Avoid:", 0, 1)

    pdf.set_font("DejaVu", "", 10)
    pdf.set_text_color(40, 40, 40)
    eff_w = pdf.w - pdf.l_margin - pdf.r_margin
    for item in TP_FINAL_NOTES["avoid"]:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(eff_w, 5.5, f"  * {item}", 0, "L")

    pdf.ln(3)
    pdf.set_font("DejaVu", "B", 11)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 7, "Pre-Interview Checklist:", 0, 1)

    pdf.set_font("DejaVu", "", 10)
    pdf.set_text_color(40, 40, 40)
    for item in TP_FINAL_NOTES["checklist"]:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(eff_w, 5.5, f"  [ ] {item}", 0, "L")

    pdf.ln(5)
    pdf.set_font("DejaVu", "B", 12)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, "Good luck, Raymond! You've done the work. Now show them what you've built.", 0, 1, "C")

    pdf.output(output_path)
    print(f"  Created: {output_path}")


# ============================================================================
# MAIN
# ============================================================================

def main():
    print("Generating updated HeatShield Agri QA documents...\n")

    # QA Professional
    print("[1/4] QA Professional DOCX")
    generate_qa_professional_docx(os.path.join(DOCS_DIR, "HeatShield_QA_Professional.docx"))

    print("[2/4] QA Professional PDF")
    generate_qa_professional_pdf(os.path.join(DOCS_DIR, "HeatShield_QA_Professional.pdf"))

    print("[3/4] QA Talking Points DOCX")
    generate_qa_talking_points_docx(os.path.join(DOCS_DIR, "HeatShield_QA_Talking_Points.docx"))

    print("[4/4] QA Talking Points PDF")
    generate_qa_talking_points_pdf(os.path.join(DOCS_DIR, "HeatShield_QA_Talking_Points.pdf"))

    print("\nAll documents generated successfully!")
    print(f"\nOutput files:")
    for fname in [
        "HeatShield_QA_Professional.docx",
        "HeatShield_QA_Professional.pdf",
        "HeatShield_QA_Talking_Points.docx",
        "HeatShield_QA_Talking_Points.pdf",
    ]:
        fpath = os.path.join(DOCS_DIR, fname)
        if os.path.exists(fpath):
            size_kb = os.path.getsize(fpath) / 1024
            print(f"  [OK] {fname} ({size_kb:.1f} KB)")
        else:
            print(f"  [FAIL] {fname} - NOT FOUND")


if __name__ == "__main__":
    main()
