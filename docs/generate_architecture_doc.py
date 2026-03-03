#!/usr/bin/env python3
"""
Generate HeatShield Agri Platform Architecture Document in PDF and DOCX formats.

Dependencies:
    pip install fpdf2 python-docx

Output:
    - docs/HeatShield_Agri_Architecture.pdf
    - docs/HeatShield_Agri_Architecture.docx
"""

import os
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DOCS_DIR = Path(__file__).resolve().parent
PDF_PATH = DOCS_DIR / "HeatShield_Agri_Architecture.pdf"
DOCX_PATH = DOCS_DIR / "HeatShield_Agri_Architecture.docx"

TITLE = "HeatShield Agri \u2014 Platform Architecture Document"
SUBTITLE = "Prepared by Raymond Reuel Wayesu, Wayesu Community Research Organisation Ltd"
DATE_STR = "March 2026"
FOOTER_TEXT = "Wayesu Community Research Organisation Ltd | Raymond Reuel Wayesu"

# Colour palette (RGB tuples)
BRAND_DARK = (30, 58, 38)       # dark green
BRAND_MID = (46, 125, 50)       # medium green
BRAND_LIGHT = (200, 230, 201)   # light green tint
TABLE_HEADER_BG = (46, 125, 50)
TABLE_HEADER_FG = (255, 255, 255)
TABLE_ALT_BG = (232, 245, 233)
TABLE_BORDER = (180, 180, 180)
GREY_TEXT = (100, 100, 100)

# ---------------------------------------------------------------------------
# Document content data
# ---------------------------------------------------------------------------

SECTIONS: list[dict] = []


def _sec(number: str, title: str, paragraphs: list, table=None, bullets=None):
    """Helper to append a section."""
    SECTIONS.append({
        "number": number,
        "title": title,
        "paragraphs": paragraphs,
        "table": table,
        "bullets": bullets,
    })


# 1. Executive Summary
_sec("1", "Executive Summary", [
    "HeatShield Agri is an AI-powered heat stress early warning system designed for Uganda's "
    "12.4 million agricultural workers. The platform delivers real-time WBGT (Wet-Bulb Globe "
    "Temperature) predictions and actionable safety guidance through three complementary channels: "
    "a web dashboard, an Android app, and a USSD/SMS system for feature phones.",

    "The core scientific engine implements ISO 7243 occupational heat stress calculations "
    "consistently across all platforms. A Random Forest machine learning pipeline enhances "
    "physics-based predictions with trained ONNX models for temperature, humidity, and wind "
    "speed forecasting.",

    "This document provides a comprehensive overview of the platform architecture covering "
    "the web dashboard, Android application, USSD backend, WASM engine, and the Random Forest "
    "ML pipeline."
])

# 2. System Architecture Overview
_sec("2", "System Architecture Overview", [
    "HeatShield Agri is a multi-channel delivery platform with a shared scientific core. "
    "The architecture ensures that every channel computes WBGT identically, so that a farmer "
    "receiving guidance via USSD gets the same risk assessment as a researcher using the web dashboard."
],
    bullets=[
        "Web Dashboard (React + TypeScript) \u2014 production deployed at heatshieldagri.vercel.app",
        "Android App (Kotlin + Jetpack Compose) \u2014 functional prototype with ML inference",
        "USSD Backend (Rust + Actix-web) \u2014 structured skeleton with full menu system",
        "WASM Engine (Rust) \u2014 complete calculation library for browser-native performance",
        "Random Forest ML Pipeline \u2014 ONNX models deployed on web and Android",
    ]
)

_sec("2.1", "Shared Scientific Foundation", [
    "All platforms share the following scientific components to guarantee cross-platform consistency:"
],
    bullets=[
        "ISO 7243 WBGT formula: WBGT = 0.7*Tw + 0.2*Tg + 0.1*Ta",
        "Stull (2011) wet-bulb temperature approximation",
        "Liljegren globe temperature (Newton-Raphson iterative solver)",
        "Risk thresholds: Low (<26\u00b0C), Moderate (26\u201328\u00b0C), High (28\u201330\u00b0C), Very High (30\u201332\u00b0C), Extreme (>32\u00b0C)",
        "Open-Meteo weather API for real-time data",
        "12 Uganda districts with GPS coordinates",
    ]
)

# 3. Web Dashboard Architecture
_sec("3", "Web Dashboard Architecture", [
    "The web dashboard is the primary interface for researchers, extension officers, and "
    "government officials. It is a single-page application built with modern web technologies "
    "and deployed to production on Vercel."
],
    table={
        "headers": ["Attribute", "Detail"],
        "rows": [
            ["Stack", "React 18.2, TypeScript, Vite, TailwindCSS 3.4, Zustand, React Query, Recharts"],
            ["Deployment", "Vercel (production, with security headers)"],
            ["Live URL", "https://heatshieldagri.vercel.app"],
        ]
    }
)

_sec("3.1", "Web \u2014 Key Components", [
    "The web dashboard is composed of several feature modules, each addressing a specific "
    "user need:"
],
    bullets=[
        "Dashboard: Real-time WBGT with current conditions for the selected district.",
        "Forecast: 72-hour WBGT forecast with hourly granularity; BOTH Physics and Random Forest models displayed side by side.",
        "Heat Map: Interactive SVG map with IDW (Inverse Distance Weighting) spatial interpolation across all 12 districts.",
        "Schedule Optimizer: AI-powered work schedule recommendations based on WBGT forecast and daylight hours (06:00\u201318:00).",
        "Alerts: SMS alert registration with configurable WBGT thresholds and multi-language support.",
        "Demo: Interactive WBGT calculator and WASM compilation demonstration.",
        "Multi-language: English, Luganda, Runyankole, Acholi, Swahili.",
    ]
)

_sec("3.2", "Web \u2014 ML Integration", [
    "The web dashboard integrates the Random Forest ML pipeline directly in the browser "
    "using ONNX Runtime Web for native inference."
],
    bullets=[
        "3 Random Forest ONNX models (~6.2 MB total): temperature_model.onnx, humidity_model.onnx, windspeed_model.onnx.",
        "17-feature engineering pipeline: 8 lag features (1, 2, 3, 6, 12, 24, 48, 72 h), 4 cyclical time encodings (hour sin/cos, day-of-year sin/cos), 3 rolling statistics (mean_24h, mean_72h, std_24h), 2 delta features.",
        "Physics-first architecture: NWP forecast \u2192 ISO 7243 baseline \u2192 RF validation (MAE \u2264 2.0\u00b0C threshold) \u2192 70% physics + 30% RF blend.",
        "ONNX Runtime Web for browser-native inference with zero server dependency.",
    ]
)

# 4. Android App Architecture
_sec("4", "Android App Architecture", [
    "The Android application brings HeatShield Agri to mobile users in the field. "
    "Built with Kotlin and Jetpack Compose, it follows the MVVM architecture pattern "
    "with Hilt dependency injection."
],
    table={
        "headers": ["Attribute", "Detail"],
        "rows": [
            ["Stack", "Kotlin, Jetpack Compose, Material 3, Hilt DI, Coroutines, StateFlow, Retrofit, ONNX Runtime Android"],
            ["Min SDK", "26 (Android 8.0)"],
            ["Target SDK", "34"],
            ["Architecture", "MVVM with Hilt dependency injection"],
            ["Branding", "\"Designed by Wayesu Community Research Organisation Ltd\" in Dashboard footer"],
        ]
    }
)

_sec("4.1", "Android \u2014 Key Components", [
    "The Android app mirrors the web dashboard's feature set with native mobile "
    "optimisations:"
],
    bullets=[
        "Dashboard Screen: Real-time weather from Open-Meteo with live WBGT calculation.",
        "Forecast Screen: Dual-model display showing BOTH Physics and Random Forest WBGT side by side for every hour, with district selector, daylight hours (06:00\u201318:00), and summary comparison cards.",
        "Schedule Screen: ML-powered work schedule optimizer with district selector, work hours selector, and productivity scoring.",
        "Map Screen: Heat risk visualisation across Uganda districts.",
        "Demo Screen: Interactive WBGT calculator with sliders for temperature, humidity, and wind speed.",
        "Alerts Screen: Alert configuration UI for threshold-based notifications.",
    ]
)

_sec("4.2", "Android \u2014 Architecture Components", [
    "The app follows clean architecture principles with clear separation of concerns:"
],
    bullets=[
        "ViewModels: ForecastViewModel, ScheduleViewModel (both implement physics-first + RF validation).",
        "Repository: WeatherRepository with getWeatherWithHistory() for historical + forecast data.",
        "ML Engine: HeatShieldMLInference singleton (ONNX Runtime with 2 inference threads).",
        "DI Module: MLModule provides singleton ML inference instance via Hilt.",
        "API: WeatherApiService with Retrofit (Open-Meteo, past_days=4 for 96 h history).",
    ]
)

_sec("4.3", "Android \u2014 ML Integration", [
    "The Android ML pipeline is identical to the web implementation, ensuring cross-platform "
    "consistency:"
],
    bullets=[
        "3 ONNX models in assets/models/ (~6.2 MB total).",
        "ONNX Runtime Android (native C++ via JNI) for high-performance inference.",
        "17-feature pipeline identical to web (lag, cyclical, rolling, delta features).",
        "Physics-first architecture with RF validation and blending.",
    ]
)

# 5. USSD Backend Architecture
_sec("5", "USSD Backend Architecture", [
    "The USSD backend enables HeatShield Agri to reach the majority of Ugandan farmers who "
    "rely on feature phones without internet access. Built in Rust for maximum reliability "
    "and minimal resource usage."
],
    table={
        "headers": ["Attribute", "Detail"],
        "rows": [
            ["Stack", "Rust 2021, Actix-web 4, Tokio, SQLx (PostgreSQL), Redis, Africa's Talking integration"],
            ["Binary Size", "6.3 MB statically linked ELF"],
            ["Service Code", "*384*HEAT#"],
        ]
    }
)

_sec("5.1", "USSD \u2014 Key Components", [
    "The USSD system provides a text-based menu interface for feature phone users:"
],
    bullets=[
        "USSD Callback Handler: Parses Africa's Talking format (sessionId, phoneNumber, text, serviceCode).",
        "Menu System: 6-item main menu \u2014 (1) Check Today's Heat Risk, (2) 3-Day Forecast, (3) Safe Work Hours, (4) Register for SMS Alerts, (5) Change Location, (6) Change Language.",
        "Session Management: State machine with enum states (MainMenu, ViewingForecast, RegisteringLocation, etc.).",
        "i18n: 5 languages (English, Luganda, Runyankole, Acholi, Swahili) with 15+ translated message keys.",
        "WBGT Engine: Same ISO 7243 calculations implemented natively in Rust.",
        "Location Service: 16 Uganda districts, haversine distance calculation, region filtering.",
        "Configuration: Environment variables for database, Redis, API keys, session timeout.",
    ]
)

_sec("5.2", "USSD \u2014 Infrastructure", [
    "The following infrastructure components are configured but not yet connected to live services:"
],
    bullets=[
        "PostgreSQL for user data persistence.",
        "Redis for session management.",
        "Africa's Talking for SMS delivery.",
        "Docker multi-stage build (rust:1.74-slim builder, debian:bookworm-slim runtime).",
    ]
)

# 6. WASM Engine Architecture
_sec("6", "WASM Engine Architecture", [
    "The WASM engine is a Rust library that compiles to WebAssembly, providing browser-native "
    "performance for computationally intensive heat stress calculations. The complete source "
    "code is ready; a JavaScript fallback is currently used in production while WASM "
    "compilation is finalised."
],
    table={
        "headers": ["Module", "Description", "Lines of Code"],
        "rows": [
            ["wbgt.rs", "ISO 7243 WBGT calculation", "213"],
            ["risk.rs", "Risk classification with local language messages", "250"],
            ["spatial.rs", "IDW spatial interpolation, 5 km grid for Uganda", "244"],
            ["schedule.rs", "Work schedule optimisation", "275"],
        ]
    }
)

_sec("6.1", "WASM \u2014 Technical Details", [
    "The WASM engine is built with wasm-bindgen, js-sys, and web-sys for seamless JavaScript "
    "interoperability. Each module includes comprehensive unit tests."
],
    bullets=[
        "Stack: Rust, wasm-bindgen, js-sys, web-sys.",
        "Each module has its own unit test suite.",
        "Status: Complete source code; not yet compiled to WASM (JavaScript fallback used in production).",
    ]
)

# 7. Random Forest ML Pipeline
_sec("7", "Random Forest ML Pipeline", [
    "The ML pipeline enhances the physics-based WBGT predictions by training Random Forest "
    "models on historical weather data. The models are exported to ONNX format for "
    "cross-platform deployment."
],
    table={
        "headers": ["Attribute", "Detail"],
        "rows": [
            ["Model Type", "Random Forest (scikit-learn trained, exported to ONNX)"],
            ["Number of Models", "3 (temperature, humidity, wind speed)"],
            ["Training Data", "Historical Open-Meteo weather data"],
            ["Feature Count", "17 features per prediction"],
            ["Deployment (Web)", "ONNX Runtime Web (JavaScript)"],
            ["Deployment (Android)", "ONNX Runtime Android (native C++ via JNI)"],
            ["Architecture", "Physics-first: NWP baseline; RF validates when MAE <= 2.0\u00b0C"],
        ]
    }
)

_sec("7.1", "Feature Engineering", [
    "Each prediction uses a 17-feature vector engineered from historical weather observations:"
],
    table={
        "headers": ["Category", "Features", "Count"],
        "rows": [
            ["Lag features", "t-1, t-2, t-3, t-6, t-12, t-24, t-48, t-72 hours", "8"],
            ["Cyclical encodings", "sin/cos of hour-of-day, sin/cos of day-of-year", "4"],
            ["Rolling statistics", "mean_24h, mean_72h, std_24h", "3"],
            ["Delta features", "delta_1h, delta_24h", "2"],
        ]
    }
)

_sec("7.2", "Inference Architecture", [
    "The physics-first architecture ensures that the system always has a reliable baseline, "
    "with ML enhancing predictions only when it demonstrates sufficient accuracy:"
],
    bullets=[
        "Step 1: Fetch NWP (Numerical Weather Prediction) forecast from Open-Meteo.",
        "Step 2: Compute physics WBGT using ISO 7243.",
        "Step 3: If RF models available, run the 17-feature pipeline to predict temperature, humidity, and wind speed.",
        "Step 4: Compute RF WBGT from the ML-predicted weather variables.",
        "Step 5: Compare RF vs Physics using MAE (Mean Absolute Error).",
        "Step 6: If MAE <= 2.0\u00b0C, blend results (70% physics + 30% RF); otherwise, use physics only.",
    ]
)

# 8. Cross-Platform Consistency
_sec("8", "Cross-Platform Consistency", [
    "A core design principle of HeatShield Agri is that every platform must produce identical "
    "results for the same inputs. The following table confirms that all platforms share the "
    "same scientific and data foundations."
],
    table={
        "headers": ["Component", "Web", "Android", "USSD", "WASM"],
        "rows": [
            ["WBGT Formula", "ISO 7243", "ISO 7243", "ISO 7243", "ISO 7243"],
            ["Wet-bulb Method", "Stull (2011)", "Stull (2011)", "Stull (2011)", "Stull (2011)"],
            ["Globe Temp Method", "Liljegren (N-R)", "Liljegren (N-R)", "Liljegren (N-R)", "Liljegren (N-R)"],
            ["Risk Thresholds", "5 levels", "5 levels", "5 levels", "5 levels"],
            ["Weather API", "Open-Meteo", "Open-Meteo", "Open-Meteo", "N/A (lib)"],
            ["Languages", "5", "5", "5", "5"],
        ]
    }
)

# 9. Data Flow
_sec("9", "Data Flow", [
    "All platforms consume weather data from the Open-Meteo API, which aggregates forecasts "
    "from ECMWF and GFS numerical weather prediction models."
],
    table={
        "headers": ["Parameter", "Detail"],
        "rows": [
            ["Weather Source", "Open-Meteo API (ECMWF + GFS models)"],
            ["Current Variables", "temperature_2m, relative_humidity_2m, wind_speed_10m, weather_code, shortwave_radiation"],
            ["Historical Window", "past_days=4, forecast_days=2 (96 h history + 48 h forecast)"],
            ["Timezone", "Africa/Kampala"],
        ]
    }
)

_sec("9.1", "Processing Pipeline", [
    "The data processing pipeline follows the same sequence across all platforms:"
],
    bullets=[
        "Step 1: Fetch weather data from Open-Meteo API.",
        "Step 2: Compute physics WBGT using ISO 7243 (Stull wet-bulb, Liljegren globe temperature).",
        "Step 3: If RF models are available, run the 17-feature engineering pipeline.",
        "Step 4: Predict temperature, humidity, and wind speed via ONNX inference.",
        "Step 5: Compute RF WBGT from predicted variables.",
        "Step 6: Compare RF vs Physics (MAE); display both models or blend if within threshold.",
    ]
)

# 10. Security and Deployment
_sec("10", "Security and Deployment", [
    "Each platform follows security best practices appropriate to its deployment context."
],
    table={
        "headers": ["Platform", "Security & Deployment Measures"],
        "rows": [
            ["Web", "Vercel with security headers, SPA routing, asset caching, HTTPS, Content-Security-Policy"],
            ["Android", "Hilt DI, ProGuard minification, release signing, secure API communication"],
            ["USSD", "Docker containerisation, non-root user, TLS, environment variable secrets management"],
            ["Data Policy", "Minimal collection (phone number + district only), TLS encryption, no data sales"],
        ]
    }
)


# ============================================================================
# PDF GENERATION
# ============================================================================

def generate_pdf():
    """Generate the architecture document as a professional PDF using fpdf2."""
    from fpdf import FPDF

    class ArchPDF(FPDF):
        FONT = "DejaVu"

        def __init__(self):
            super().__init__(orientation="P", unit="mm", format="A4")
            self.set_auto_page_break(auto=True, margin=25)
            self._section_title_for_header = ""
            # Register DejaVu Sans (Unicode-capable)
            font_dir = "/usr/share/fonts/truetype/dejavu"
            self.add_font("DejaVu", "", f"{font_dir}/DejaVuSans.ttf")
            self.add_font("DejaVu", "B", f"{font_dir}/DejaVuSans-Bold.ttf")
            # No italic TTF available for DejaVu Sans, use Book as italic substitute
            self.add_font("DejaVu", "I", f"{font_dir}/DejaVuSans.ttf")
            self.add_font("DejaVu", "BI", f"{font_dir}/DejaVuSans-Bold.ttf")

        def header(self):
            if self.page_no() == 1:
                return  # Cover page has custom header
            # Thin green line at top
            self.set_draw_color(*BRAND_MID)
            self.set_line_width(0.6)
            self.line(10, 10, 200, 10)
            # Header text
            self.set_font(self.FONT, "I", 8)
            self.set_text_color(*GREY_TEXT)
            self.set_y(12)
            self.cell(0, 5, "HeatShield Agri \u2014 Platform Architecture Document", align="L")
            self.cell(0, 5, f"Page {self.page_no()}", align="R")
            self.ln(8)

        def footer(self):
            self.set_y(-18)
            self.set_draw_color(*BRAND_MID)
            self.set_line_width(0.4)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(2)
            self.set_font(self.FONT, "I", 7)
            self.set_text_color(*GREY_TEXT)
            self.cell(0, 5, FOOTER_TEXT, align="C")

        def cover_page(self):
            self.add_page()
            lm = self.l_margin
            pw = self.w - self.l_margin - self.r_margin  # printable width

            # Green header band
            self.set_fill_color(*BRAND_DARK)
            self.rect(0, 0, 210, 90, "F")

            # Title
            self.set_xy(lm, 25)
            self.set_font(self.FONT, "B", 28)
            self.set_text_color(255, 255, 255)
            self.cell(pw, 12, "HeatShield Agri", align="C")
            self.set_xy(lm, 42)
            self.set_font(self.FONT, "", 16)
            self.cell(pw, 10, "Platform Architecture Document", align="C")

            # Accent line
            self.set_draw_color(200, 230, 201)
            self.set_line_width(1.2)
            self.line(60, 72, 150, 72)

            # Subtitle block
            self.set_xy(lm, 100)
            self.set_font(self.FONT, "", 12)
            self.set_text_color(*BRAND_DARK)
            self.cell(pw, 8, "Prepared by", align="C")

            self.set_xy(lm, 112)
            self.set_font(self.FONT, "B", 14)
            self.cell(pw, 8, "Raymond Reuel Wayesu", align="C")

            self.set_xy(lm, 124)
            self.set_font(self.FONT, "", 11)
            self.set_text_color(*GREY_TEXT)
            self.cell(pw, 8, "Wayesu Community Research Organisation Ltd", align="C")

            self.set_xy(lm, 140)
            self.set_font(self.FONT, "", 11)
            self.set_text_color(*BRAND_DARK)
            self.cell(pw, 8, DATE_STR, align="C")

            # Decorative box
            self.set_fill_color(*BRAND_LIGHT)
            self.set_draw_color(*BRAND_MID)
            self.set_line_width(0.4)
            self.rect(30, 165, 150, 40, "DF")
            self.set_xy(30, 170)
            self.set_font(self.FONT, "I", 10)
            self.set_text_color(*BRAND_DARK)
            self.multi_cell(150, 6,
                "An AI-powered heat stress early warning system\n"
                "for Uganda's 12.4 million agricultural workers.\n"
                "Delivering real-time WBGT predictions via Web, Android, and USSD.",
                align="C"
            )

            # Classification
            self.set_xy(lm, 230)
            self.set_font(self.FONT, "B", 9)
            self.set_text_color(*GREY_TEXT)
            self.cell(pw, 6, "CONFIDENTIAL", align="C")

        def table_of_contents(self):
            self.add_page()
            self.set_font(self.FONT, "B", 20)
            self.set_text_color(*BRAND_DARK)
            self.cell(0, 12, "Table of Contents", align="L")
            self.ln(15)

            for sec in SECTIONS:
                num = sec["number"]
                title = sec["title"]
                is_sub = "." in num
                indent = 10 if is_sub else 0
                font_style = "" if is_sub else "B"
                font_size = 10 if is_sub else 11

                self.set_x(15 + indent)
                self.set_font(self.FONT, font_style, font_size)
                self.set_text_color(*BRAND_DARK)
                self.cell(12, 7, num)
                self.cell(0, 7, title)
                self.ln(7 if is_sub else 8)

        def section_heading(self, number: str, title: str):
            is_sub = "." in number
            if not is_sub:
                # Major section: green bar
                self.ln(6)
                self.set_fill_color(*BRAND_MID)
                self.set_text_color(255, 255, 255)
                self.set_font(self.FONT, "B", 14)
                self.cell(0, 10, f"  {number}.  {title}", fill=True)
                self.ln(12)
            else:
                self.ln(4)
                self.set_text_color(*BRAND_DARK)
                self.set_font(self.FONT, "B", 12)
                self.cell(0, 8, f"{number}  {title}")
                self.ln(10)

        def body_text(self, text: str):
            self.set_font(self.FONT, "", 10)
            self.set_text_color(40, 40, 40)
            self.multi_cell(0, 5.5, text)
            self.ln(3)

        def bullet_list(self, items: list[str]):
            for item in items:
                y0 = self.get_y()
                # Check space
                if y0 > 260:
                    self.add_page()
                # Bullet character
                self.set_x(18)
                self.set_font(self.FONT, "B", 10)
                self.set_text_color(*BRAND_MID)
                self.cell(5, 5.5, "\u2022")
                # Item text
                self.set_font(self.FONT, "", 10)
                self.set_text_color(40, 40, 40)
                self.multi_cell(162, 5.5, item)
                self.ln(1.5)

        def data_table(self, headers: list[str], rows: list[list[str]]):
            # Calculate column widths
            n_cols = len(headers)
            avail = 185
            if n_cols == 2:
                col_w = [55, 130]
            elif n_cols == 3:
                col_w = [50, 95, 40]
            elif n_cols == 5:
                col_w = [40, 35, 35, 35, 35]
            else:
                w = avail // n_cols
                col_w = [w] * n_cols
                col_w[-1] = avail - w * (n_cols - 1)

            row_h = 7

            def _draw_row(cell_values, is_header=False, alt=False):
                # Determine row height needed
                max_lines = 1
                for i, cv in enumerate(cell_values):
                    self.set_font(self.FONT, "B" if is_header else "", 9)
                    cw = col_w[i] - 2
                    n_lines = max(1, len(self.multi_cell(cw, row_h, cv, dry_run=True, output="LINES")))
                    max_lines = max(max_lines, n_lines)
                h = row_h * max_lines

                # Check page break
                if self.get_y() + h > 270:
                    self.add_page()

                y0 = self.get_y()
                x0 = 12.5

                for i, cv in enumerate(cell_values):
                    x = x0 + sum(col_w[:i])
                    self.set_xy(x, y0)
                    if is_header:
                        self.set_fill_color(*TABLE_HEADER_BG)
                        self.set_text_color(*TABLE_HEADER_FG)
                        self.set_font(self.FONT, "B", 9)
                    else:
                        if alt:
                            self.set_fill_color(*TABLE_ALT_BG)
                        else:
                            self.set_fill_color(255, 255, 255)
                        self.set_text_color(40, 40, 40)
                        self.set_font(self.FONT, "", 9)

                    # Draw cell background
                    self.rect(x, y0, col_w[i], h, "F")
                    self.set_draw_color(*TABLE_BORDER)
                    self.set_line_width(0.2)
                    self.rect(x, y0, col_w[i], h, "D")

                    # Write text
                    self.set_xy(x + 1, y0 + 1)
                    self.multi_cell(col_w[i] - 2, row_h, cv)

                self.set_y(y0 + h)

            _draw_row(headers, is_header=True)
            for idx, row in enumerate(rows):
                _draw_row(row, alt=(idx % 2 == 1))
            self.ln(5)

    # Build the PDF
    pdf = ArchPDF()
    pdf.set_title(TITLE)
    pdf.set_author("Raymond Reuel Wayesu")
    pdf.set_subject("HeatShield Agri Platform Architecture")

    # Cover page
    pdf.cover_page()

    # Table of contents
    pdf.table_of_contents()

    # Content sections
    for sec in SECTIONS:
        num = sec["number"]
        is_major = "." not in num
        if is_major:
            pdf.add_page()

        pdf.section_heading(num, sec["title"])

        for para in sec["paragraphs"]:
            pdf.body_text(para)

        if sec.get("table"):
            pdf.data_table(sec["table"]["headers"], sec["table"]["rows"])

        if sec.get("bullets"):
            pdf.bullet_list(sec["bullets"])

    # Save
    pdf.output(str(PDF_PATH))
    print(f"[OK] PDF generated: {PDF_PATH}")
    print(f"     Size: {PDF_PATH.stat().st_size:,} bytes")


# ============================================================================
# DOCX GENERATION
# ============================================================================

def generate_docx():
    """Generate the architecture document as a professional DOCX using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # ---- Styles setup ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(40, 40, 40)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def add_footer(section):
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(FOOTER_TEXT)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.italic = True

    add_footer(doc.sections[0])

    # ---- Cover Page ----
    for _ in range(4):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HeatShield Agri")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(*BRAND_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Platform Architecture Document")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(*BRAND_MID)

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 40)
    run.font.color.rgb = RGBColor(*BRAND_MID)
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Author info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Raymond Reuel Wayesu")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(*BRAND_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Wayesu Community Research Organisation Ltd")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(DATE_STR)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*BRAND_DARK)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tagline box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "An AI-powered heat stress early warning system for Uganda's 12.4 million "
        "agricultural workers. Delivering real-time WBGT predictions via Web, Android, and USSD."
    )
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(*BRAND_DARK)

    # Page break
    doc.add_page_break()

    # ---- Table of Contents ----
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(*BRAND_DARK)
    p.space_after = Pt(18)

    for sec in SECTIONS:
        num = sec["number"]
        title = sec["title"]
        is_sub = "." in num
        p = doc.add_paragraph()
        if is_sub:
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(f"{num}  {title}")
            run.font.size = Pt(11)
        else:
            run = p.add_run(f"{num}.  {title}")
            run.bold = True
            run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(*BRAND_DARK)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ---- Helper functions ----
    def add_section_heading(number, title):
        is_sub = "." in number
        if is_sub:
            p = doc.add_heading(f"{number}  {title}", level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(*BRAND_DARK)
                run.font.size = Pt(14)
        else:
            p = doc.add_heading(f"{number}.  {title}", level=1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(*BRAND_MID)
                run.font.size = Pt(18)
            # Add thin green line via border bottom
            pf = p.paragraph_format
            pf.space_before = Pt(18)
            pf.space_after = Pt(10)

    def add_body_paragraph(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(11)
        return p

    def add_bullet_list(items):
        for item in items:
            p = doc.add_paragraph(item, style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.size = Pt(11)

    def set_cell_shading(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def add_data_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        hdr = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, "2E7D32")

        # Data rows
        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx + 1]
            for c_idx, val in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(10)
                if r_idx % 2 == 1:
                    set_cell_shading(cell, "E8F5E9")

        # Set column widths for readability
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.space_before = Pt(2)

        doc.add_paragraph()  # spacing after table

    # ---- Generate content sections ----
    for sec in SECTIONS:
        num = sec["number"]
        is_major = "." not in num
        if is_major:
            doc.add_page_break()

        add_section_heading(num, sec["title"])

        for para in sec["paragraphs"]:
            add_body_paragraph(para)

        if sec.get("table"):
            add_data_table(sec["table"]["headers"], sec["table"]["rows"])

        if sec.get("bullets"):
            add_bullet_list(sec["bullets"])

    # Save
    doc.save(str(DOCX_PATH))
    print(f"[OK] DOCX generated: {DOCX_PATH}")
    print(f"     Size: {DOCX_PATH.stat().st_size:,} bytes")


# ============================================================================
# Main
# ============================================================================

def main():
    print("=" * 60)
    print("HeatShield Agri — Architecture Document Generator")
    print("=" * 60)
    print()

    generate_pdf()
    print()
    generate_docx()

    print()
    print("-" * 60)
    print("Generation complete. Output files:")
    print(f"  PDF:  {PDF_PATH}")
    print(f"  DOCX: {DOCX_PATH}")
    print("-" * 60)


if __name__ == "__main__":
    main()
