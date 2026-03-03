#!/usr/bin/env python3
"""
Generate PDF and DOCX versions of the Deployment Readiness Strategy Report.
Reads the updated Markdown file and produces formatted output documents.

Author: Raymond Reuel Wayesu
Organisation: Wayesu Community Research Organisation Ltd
"""

import re
import os
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

MD_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "Deployment_Readiness_Strategy_Report.md")
PDF_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "Deployment_Readiness_Strategy_Report.pdf")
DOCX_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                         "Deployment_Readiness_Strategy_Report.docx")

FOOTER_TEXT = "Wayesu Community Research Organisation Ltd | Raymond Reuel Wayesu"

# DejaVu font paths (Unicode-capable)
DEJAVU_DIR = "/usr/share/fonts/truetype/dejavu"
DEJAVU_SANS = os.path.join(DEJAVU_DIR, "DejaVuSans.ttf")
DEJAVU_SANS_BOLD = os.path.join(DEJAVU_DIR, "DejaVuSans-Bold.ttf")
DEJAVU_MONO = os.path.join(DEJAVU_DIR, "DejaVuSansMono.ttf")
DEJAVU_MONO_BOLD = os.path.join(DEJAVU_DIR, "DejaVuSansMono-Bold.ttf")


# ---------------------------------------------------------------------------
# Markdown parser helpers
# ---------------------------------------------------------------------------

def read_markdown(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def parse_table(lines):
    """Parse a markdown table into a list of rows (each row is a list of cell strings)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        # Skip separator rows like |---|---|---|
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        rows.append(cells)
    return rows


def parse_md_blocks(md_text):
    """
    Parse the markdown into a list of block dicts:
      { 'type': 'h1'|'h2'|'h3'|'h4'|'para'|'bullet'|'table'|'hr',
        'content': str or list }
    """
    lines = md_text.split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Horizontal rule
        if stripped == "---":
            blocks.append({"type": "hr", "content": ""})
            i += 1
            continue

        # Headers
        if stripped.startswith("#### "):
            blocks.append({"type": "h4", "content": stripped[5:]})
            i += 1
            continue
        if stripped.startswith("### "):
            blocks.append({"type": "h3", "content": stripped[4:]})
            i += 1
            continue
        if stripped.startswith("## "):
            blocks.append({"type": "h2", "content": stripped[3:]})
            i += 1
            continue
        if stripped.startswith("# "):
            blocks.append({"type": "h1", "content": stripped[2:]})
            i += 1
            continue

        # Table (collect consecutive lines starting with |)
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if rows:
                blocks.append({"type": "table", "content": rows})
            continue

        # Bullet / numbered list
        if re.match(r'^(\d+\.\s+|- )', stripped):
            blocks.append({"type": "bullet", "content": stripped})
            i += 1
            continue

        # Non-empty paragraph
        if stripped:
            blocks.append({"type": "para", "content": stripped})
            i += 1
            continue

        # Empty line
        i += 1

    return blocks


def strip_md_formatting(text):
    """Remove markdown bold/italic markers for plain-text contexts."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


# ---------------------------------------------------------------------------
# PDF generation with fpdf2
# ---------------------------------------------------------------------------

class ReportPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=20)
        # Register Unicode-capable fonts
        self.add_font("DejaVu", "", DEJAVU_SANS)
        self.add_font("DejaVu", "B", DEJAVU_SANS_BOLD)
        self.add_font("DejaVuMono", "", DEJAVU_MONO)
        self.add_font("DejaVuMono", "B", DEJAVU_MONO_BOLD)

    def header(self):
        self.set_font("DejaVu", "B", 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 6, "HeatShield Agri - Deployment Readiness Strategy Report",
                  new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.line(10, self.get_y(), self.w - 10, self.get_y())
        self.ln(3)

    def footer(self):
        self.set_y(-15)
        self.set_font("DejaVu", "", 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 10, FOOTER_TEXT + f"  |  Page {self.page_no()}/{{nb}}",
                  align="C", new_x=XPos.RIGHT, new_y=YPos.TOP)

    def add_heading(self, text, level=1):
        sizes = {1: 18, 2: 14, 3: 12, 4: 10}
        size = sizes.get(level, 10)
        self.set_font("DejaVu", "B", size)
        if level <= 2:
            self.set_text_color(0, 90, 50)
        else:
            self.set_text_color(0, 60, 30)
        self.ln(3)
        clean = strip_md_formatting(text)
        self.multi_cell(0, size * 0.55, clean)
        self.set_text_color(0, 0, 0)
        self.ln(2)

    def add_paragraph(self, text):
        self.set_font("DejaVu", "", 9)
        self.set_text_color(0, 0, 0)
        self._write_rich_text(text, 9)
        self.ln(4)

    def add_bullet(self, text):
        self.set_font("DejaVu", "", 9)
        self.set_text_color(0, 0, 0)
        clean = re.sub(r'^(\d+\.\s+|- )', '', text.strip())
        marker = text.strip()
        if marker.startswith("-"):
            marker_str = "  \u2022  "
        else:
            m = re.match(r'^(\d+\.)', marker)
            marker_str = f"  {m.group(1)} " if m else "  \u2022  "

        x_start = self.get_x()
        self.set_x(x_start + 5)
        self.set_font("DejaVu", "B", 9)
        self.write(5, marker_str)
        self._write_rich_text(clean, 9)
        self.ln(5)

    def _write_rich_text(self, text, size):
        """Write text with **bold** support."""
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                self.set_font("DejaVu", "B", size)
                self.write(5, part[2:-2])
                self.set_font("DejaVu", "", size)
            else:
                subparts = re.split(r'(`.*?`)', part)
                for sp in subparts:
                    if sp.startswith("`") and sp.endswith("`"):
                        self.set_font("DejaVuMono", "", size - 1)
                        self.write(5, sp[1:-1])
                        self.set_font("DejaVu", "", size)
                    else:
                        self.write(5, sp)

    def add_table(self, rows):
        if not rows:
            return
        num_cols = len(rows[0])
        available_width = self.w - 20  # margins

        # Determine column widths based on content
        col_widths = []
        for c in range(num_cols):
            max_len = max(len(str(rows[r][c])) if c < len(rows[r]) else 0
                         for r in range(len(rows)))
            col_widths.append(max(max_len, 5))

        total = sum(col_widths)
        col_widths = [w / total * available_width for w in col_widths]

        # Cap any single column at 55% of available width
        for idx in range(len(col_widths)):
            if col_widths[idx] > available_width * 0.55:
                col_widths[idx] = available_width * 0.55

        # Re-normalize
        total = sum(col_widths)
        col_widths = [w / total * available_width for w in col_widths]

        line_height = 5.0
        self.set_font("DejaVu", "", 7)

        for r_idx, row in enumerate(rows):
            # Check page break
            if self.get_y() + 12 > self.h - 20:
                self.add_page()

            if r_idx == 0:
                self.set_font("DejaVu", "B", 7)
                self.set_fill_color(0, 90, 50)
                self.set_text_color(255, 255, 255)
            else:
                self.set_font("DejaVu", "", 7)
                if r_idx % 2 == 0:
                    self.set_fill_color(240, 248, 240)
                else:
                    self.set_fill_color(255, 255, 255)
                self.set_text_color(0, 0, 0)

            # Calculate row height needed
            max_lines = 1
            for c_idx, cell in enumerate(row):
                if c_idx < len(col_widths):
                    cw = col_widths[c_idx]
                    cell_text = strip_md_formatting(str(cell))
                    char_width = 1.6
                    chars_per_line = max(int(cw / char_width), 10)
                    lines_needed = max(1, (len(cell_text) + chars_per_line - 1) // chars_per_line)
                    max_lines = max(max_lines, lines_needed)

            row_height = line_height * max_lines

            x_start = self.get_x()
            y_start = self.get_y()

            for c_idx, cell in enumerate(row):
                if c_idx >= len(col_widths):
                    break
                cw = col_widths[c_idx]
                cell_text = strip_md_formatting(str(cell))

                self.set_xy(x_start + sum(col_widths[:c_idx]), y_start)
                self.multi_cell(cw, line_height, cell_text, border=1,
                                fill=True, new_x=XPos.RIGHT, new_y=YPos.TOP)

            self.set_xy(x_start, y_start + row_height)

        self.set_text_color(0, 0, 0)
        self.ln(4)

    def add_hr(self):
        y = self.get_y()
        self.set_draw_color(200, 200, 200)
        self.line(10, y, self.w - 10, y)
        self.ln(4)


def generate_pdf(blocks, output_path):
    pdf = ReportPDF()
    pdf.alias_nb_pages()
    pdf.add_page()

    for block in blocks:
        btype = block["type"]
        content = block["content"]

        if btype == "h1":
            pdf.add_heading(content, 1)
        elif btype == "h2":
            pdf.add_heading(content, 2)
        elif btype == "h3":
            pdf.add_heading(content, 3)
        elif btype == "h4":
            pdf.add_heading(content, 4)
        elif btype == "para":
            pdf.add_paragraph(content)
        elif btype == "bullet":
            pdf.add_bullet(content)
        elif btype == "table":
            pdf.add_table(content)
        elif btype == "hr":
            pdf.add_hr()

    pdf.output(output_path)
    print(f"PDF generated: {output_path}")


# ---------------------------------------------------------------------------
# DOCX generation with python-docx
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_rich_text_to_paragraph(paragraph, text, bold_default=False, size=10):
    """Add text to a paragraph, handling **bold** markers."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(size)
        else:
            subparts = re.split(r'(`.*?`)', part)
            for sp in subparts:
                if sp.startswith("`") and sp.endswith("`"):
                    run = paragraph.add_run(sp[1:-1])
                    run.font.name = "Courier New"
                    run.font.size = Pt(size - 1)
                else:
                    run = paragraph.add_run(sp)
                    run.bold = bold_default
                    run.font.size = Pt(size)


def generate_docx(blocks, output_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(FOOTER_TEXT)
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for block in blocks:
        btype = block["type"]
        content = block["content"]

        if btype == "h1":
            p = doc.add_heading(strip_md_formatting(content), level=1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 90, 50)

        elif btype == "h2":
            p = doc.add_heading(strip_md_formatting(content), level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 90, 50)

        elif btype == "h3":
            p = doc.add_heading(strip_md_formatting(content), level=3)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 60, 30)

        elif btype == "h4":
            p = doc.add_heading(strip_md_formatting(content), level=4)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 60, 30)

        elif btype == "para":
            p = doc.add_paragraph()
            add_rich_text_to_paragraph(p, content)

        elif btype == "bullet":
            clean = re.sub(r'^(\d+\.\s+|- )', '', content.strip())
            p = doc.add_paragraph(style='List Bullet')
            add_rich_text_to_paragraph(p, clean)

        elif btype == "table":
            rows = content
            if not rows:
                continue
            num_cols = len(rows[0])
            table = doc.add_table(rows=0, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.style = 'Table Grid'

            for r_idx, row in enumerate(rows):
                docx_row = table.add_row()
                for c_idx in range(num_cols):
                    cell = docx_row.cells[c_idx]
                    cell_text = strip_md_formatting(row[c_idx]) if c_idx < len(row) else ""
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.size = Pt(9)

                    if r_idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        set_cell_shading(cell, "005A32")
                    elif r_idx % 2 == 0:
                        set_cell_shading(cell, "F0F8F0")

            doc.add_paragraph()  # spacing after table

        elif btype == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("_" * 80)
            run.font.color.rgb = RGBColor(200, 200, 200)
            run.font.size = Pt(6)

    doc.save(output_path)
    print(f"DOCX generated: {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print(f"Reading markdown from: {MD_PATH}")
    md_text = read_markdown(MD_PATH)
    blocks = parse_md_blocks(md_text)
    print(f"Parsed {len(blocks)} blocks from markdown")

    generate_pdf(blocks, PDF_PATH)
    generate_docx(blocks, DOCX_PATH)

    print("\nDone! Generated files:")
    print(f"  PDF:  {PDF_PATH}")
    print(f"  DOCX: {DOCX_PATH}")


if __name__ == "__main__":
    main()
