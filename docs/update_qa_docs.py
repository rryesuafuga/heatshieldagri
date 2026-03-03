#!/usr/bin/env python3
"""Update QA Professional and QA Talking Points documents.

Changes:
1. XGBoost → Random Forest
2. physics-informed neural network → physics-based ISO 7243
3. Raymond Wayesu → Raymond Reuel Wayesu
4. Add Wayesu Community Research Organisation Ltd
5. Update ML deployment status from planned to deployed
6. Regenerate PDFs
"""

import os
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Text replacements applied everywhere
REPLACEMENTS = [
    ("XGBoost", "Random Forest"),
    ("xgboost", "Random Forest"),
    ("physics-informed neural network", "physics-based ISO 7243 model"),
    ("Physics-informed neural network", "Physics-based ISO 7243 model"),
    ("Raymond Wayesu", "Raymond Reuel Wayesu"),
]

# More nuanced content replacements for specific sentences
SENTENCE_REPLACEMENTS = [
    (
        "XGBoost excels at handling tabular data with complex feature interactions "
        "(temperature, humidity, wind, terrain), while the physics layer ensures "
        "predictions respect known thermodynamic relationships.",
        "Random Forest excels at handling tabular data with complex feature interactions "
        "(temperature, humidity, wind, terrain) using an ensemble of decision trees, while the "
        "physics-based ISO 7243 layer ensures predictions respect known thermodynamic relationships. "
        "The Random Forest models are already deployed on Android via ONNX Runtime."
    ),
    (
        "XGBoost is computationally efficient and deployable without GPU infrastructure, "
        "critical for sustainable operations in our context. Deep learning alternatives "
        "would require expensive cloud GPU instances.",
        "Random Forest is computationally efficient and deployable without GPU infrastructure, "
        "critical for sustainable operations in our context. We've already deployed 3 ONNX models "
        "(temperature, humidity, wind speed) on Android with sub-second inference. Deep learning "
        "alternatives would require expensive cloud GPU instances."
    ),
    (
        "Lightweight by design: XGBoost and spatial interpolation, not GPU-intensive deep learning.",
        "Lightweight by design: Random Forest ensemble models and spatial interpolation, not GPU-intensive deep learning."
    ),
    (
        "Lightweight by design: We use XGBoost and spatial interpolation, not large language models or GPU-intensive deep learning.",
        "Lightweight by design: We use Random Forest ensemble models and spatial interpolation, not large language models or GPU-intensive deep learning. Our 3 ONNX models total only ~6 MB."
    ),
]


def replace_in_run(run, old, new):
    """Replace text in a single run, preserving formatting."""
    if old in run.text:
        run.text = run.text.replace(old, new)
        return True
    return False


def replace_in_paragraph(paragraph, old, new):
    """Replace text across runs in a paragraph, preserving formatting."""
    full_text = paragraph.text
    if old not in full_text:
        return False

    # Try run-level replacement first
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # If text spans multiple runs, rebuild the paragraph
    # Find which runs contain parts of the old text
    new_text = full_text.replace(old, new)

    # Clear all runs and put new text in first run
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
        return True

    return False


def replace_in_cell(cell, old, new):
    """Replace text in a table cell."""
    changed = False
    for paragraph in cell.paragraphs:
        if replace_in_paragraph(paragraph, old, new):
            changed = True
    return changed


def update_document(doc_path, extra_replacements=None):
    """Update a DOCX document with all replacements."""
    doc = Document(doc_path)
    all_replacements = SENTENCE_REPLACEMENTS + REPLACEMENTS
    if extra_replacements:
        all_replacements = extra_replacements + all_replacements

    changes = 0

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for old, new in all_replacements:
            if replace_in_paragraph(paragraph, old, new):
                changes += 1

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in all_replacements:
                    if replace_in_cell(cell, old, new):
                        changes += 1

    # Replace in headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    for old, new in all_replacements:
                        replace_in_paragraph(paragraph, old, new)
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    for old, new in all_replacements:
                        replace_in_paragraph(paragraph, old, new)

    doc.save(doc_path)
    return changes


def generate_pdf_from_docx(docx_path, pdf_path):
    """Generate PDF from DOCX using fpdf2 with DejaVu Sans font."""
    from fpdf import FPDF

    doc = Document(docx_path)

    class QAPdf(FPDF):
        def __init__(self):
            super().__init__()
            # Try to use DejaVu Sans for Unicode support
            font_dir = "/usr/share/fonts/truetype/dejavu/"
            if os.path.exists(font_dir + "DejaVuSans.ttf"):
                self.add_font("DejaVu", "", font_dir + "DejaVuSans.ttf", uni=True)
                self.add_font("DejaVu", "B", font_dir + "DejaVuSans-Bold.ttf", uni=True)
                # No oblique/italic variant available, use regular as italic
                self.add_font("DejaVu", "I", font_dir + "DejaVuSans.ttf", uni=True)
                self.font_family_name = "DejaVu"
            else:
                self.font_family_name = "Helvetica"

        def header(self):
            self.set_font(self.font_family_name, "I", 8)
            self.set_text_color(100, 100, 100)
            self.cell(0, 10, "HeatShield Agri - Interview Q&A Reference", 0, 0, "L")
            self.cell(0, 10, "Wayesu Community Research Organisation Ltd", 0, 1, "R")
            self.line(10, 18, 200, 18)
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font(self.font_family_name, "I", 8)
            self.set_text_color(100, 100, 100)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", 0, 0, "C")

    pdf = QAPdf()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    fn = pdf.font_family_name

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            pdf.ln(3)
            continue

        style = para.style.name if para.style else ""

        # Sanitize text - replace problematic characters
        text = text.replace("\u2019", "'").replace("\u2018", "'")
        text = text.replace("\u201c", '"').replace("\u201d", '"')
        text = text.replace("\u2013", "-").replace("\u2014", "--")
        text = text.replace("\u2022", "-").replace("\u2026", "...")
        text = text.replace("\u00b0", " deg ")

        if "Heading 1" in style or (style == "Title"):
            pdf.set_font(fn, "B", 16)
            pdf.set_text_color(0, 90, 50)
            pdf.multi_cell(0, 8, text)
            pdf.ln(3)
        elif "Heading 2" in style:
            pdf.set_font(fn, "B", 13)
            pdf.set_text_color(0, 100, 60)
            pdf.multi_cell(0, 7, text)
            pdf.ln(2)
        elif "Heading 3" in style or "Subtitle" in style:
            pdf.set_font(fn, "B", 11)
            pdf.set_text_color(50, 50, 50)
            pdf.multi_cell(0, 6, text)
            pdf.ln(2)
        else:
            pdf.set_font(fn, "", 10)
            pdf.set_text_color(30, 30, 30)

            # Handle bold runs
            if para.runs:
                has_bold = any(r.bold for r in para.runs if r.text.strip())
                if has_bold:
                    for run in para.runs:
                        run_text = run.text
                        run_text = run_text.replace("\u2019", "'").replace("\u2018", "'")
                        run_text = run_text.replace("\u201c", '"').replace("\u201d", '"')
                        run_text = run_text.replace("\u2013", "-").replace("\u2014", "--")
                        run_text = run_text.replace("\u2022", "-").replace("\u2026", "...")
                        run_text = run_text.replace("\u00b0", " deg ")
                        if run.bold:
                            pdf.set_font(fn, "B", 10)
                        else:
                            pdf.set_font(fn, "", 10)
                        pdf.write(5, run_text)
                    pdf.ln(5)
                    continue

            pdf.multi_cell(0, 5, text)
            pdf.ln(2)

    # Add tables
    for table in doc.tables:
        pdf.ln(3)
        col_count = len(table.columns)
        if col_count == 0:
            continue

        page_width = pdf.w - 20  # margins
        col_width = page_width / col_count

        for ri, row in enumerate(table.rows):
            # Header row
            if ri == 0:
                pdf.set_font(fn, "B", 9)
                pdf.set_fill_color(0, 90, 50)
                pdf.set_text_color(255, 255, 255)
            else:
                pdf.set_font(fn, "", 9)
                pdf.set_text_color(30, 30, 30)
                if ri % 2 == 0:
                    pdf.set_fill_color(240, 248, 240)
                else:
                    pdf.set_fill_color(255, 255, 255)

            max_h = 5
            cell_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cell_text = cell_text.replace("\u2019", "'").replace("\u2018", "'")
                cell_text = cell_text.replace("\u201c", '"').replace("\u201d", '"')
                cell_text = cell_text.replace("\u2013", "-").replace("\u2014", "--")
                cell_text = cell_text.replace("\u2022", "-").replace("\u2026", "...")
                cell_text = cell_text.replace("\u00b0", " deg ")
                cell_texts.append(cell_text)
                # Estimate height
                lines = max(1, len(cell_text) / max(1, (col_width / 2)))
                h = max(5, int(lines * 5))
                max_h = max(max_h, h)

            # Check if we need a new page
            if pdf.get_y() + max_h > pdf.h - 25:
                pdf.add_page()

            y_before = pdf.get_y()
            x_start = pdf.get_x()

            for ci, cell_text in enumerate(cell_texts):
                pdf.set_xy(x_start + ci * col_width, y_before)
                pdf.multi_cell(col_width, 5, cell_text, border=1, fill=True)

            # Move to after the tallest cell
            pdf.set_y(max(pdf.get_y(), y_before + max_h))

        pdf.ln(3)

    pdf.output(pdf_path)
    print(f"Generated PDF: {pdf_path} ({os.path.getsize(pdf_path)} bytes)")


def main():
    os.chdir("/home/user/heatshieldagri")

    # --- Update QA Professional ---
    print("=== Updating QA Professional ===")
    prof_docx = "docs/HeatShield_QA_Professional.docx"

    # Additional replacements specific to QA Professional
    prof_extra = [
        (
            "Why XGBoost with physics-informed neural network?",
            "Why Random Forest with physics-based ISO 7243 model?"
        ),
    ]

    changes = update_document(prof_docx, prof_extra)
    print(f"  {changes} replacements made in DOCX")

    # Add org name to footer area if not already there
    doc = Document(prof_docx)
    # Check if org name already in doc
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "Wayesu Community Research Organisation" not in full_text:
        # Add to end of document
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\nPrepared by Wayesu Community Research Organisation Ltd")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.italic = True
        doc.save(prof_docx)
        print("  Added org attribution")

    # Generate PDF
    generate_pdf_from_docx(prof_docx, "docs/HeatShield_QA_Professional.pdf")

    # --- Update QA Talking Points ---
    print("\n=== Updating QA Talking Points ===")
    tp_docx = "docs/HeatShield_QA_Talking_Points.docx"

    tp_extra = [
        (
            "Q: Why did you choose XGBoost combined with a physics-informed neural network?",
            "Q: Why did you choose Random Forest combined with a physics-based ISO 7243 model?"
        ),
    ]

    changes = update_document(tp_docx, tp_extra)
    print(f"  {changes} replacements made in DOCX")

    # Add org name if not present
    doc = Document(tp_docx)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "Wayesu Community Research Organisation" not in full_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\nPrepared by Wayesu Community Research Organisation Ltd")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.italic = True
        doc.save(tp_docx)
        print("  Added org attribution")

    # Generate PDF
    generate_pdf_from_docx(tp_docx, "docs/HeatShield_QA_Talking_Points.pdf")

    # --- Verify ---
    print("\n=== Verification ===")
    for path in [prof_docx, tp_docx]:
        doc = Document(path)
        remaining = 0
        for p in doc.paragraphs:
            if "XGBoost" in p.text or "xgboost" in p.text.lower():
                remaining += 1
                print(f"  WARNING: Remaining XGBoost in {path}: {p.text[:100]}")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "XGBoost" in cell.text or "xgboost" in cell.text.lower():
                        remaining += 1
                        print(f"  WARNING: Remaining XGBoost in {path} table: {cell.text[:100]}")
        if remaining == 0:
            print(f"  {path}: No XGBoost references remaining")


if __name__ == "__main__":
    main()
