"""Generate all HeatShield Agri documents: Architecture, Strategy Report, QA docs"""
import os, re
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS_DIR = '/home/user/heatshieldagri/docs'
ORG = 'Wayesu Community Research Organisation Ltd'
AUTHOR = 'Raymond Reuel Wayesu'

def s(text):
    """Sanitize text for PDF (replace unicode)"""
    reps = {'\u2014':'-','\u2013':'-','\u2018':"'",'\u2019':"'",'\u201c':'"','\u201d':'"',
            '\u2022':'*','\u2026':'...','\u00b0':'deg','\u2264':'<=','\u2265':'>=',
            '\u00d7':'x','\u2192':'->','\u2713':'[x]','\u00a9':'(c)','\u00b2':'2'}
    for o,n in reps.items(): text=text.replace(o,n)
    return text.encode('latin-1',errors='replace').decode('latin-1')

class PDF(FPDF):
    def __init__(self,title):
        super().__init__(); self.t=title; self.alias_nb_pages()
    def header(self):
        if self.page_no()>1:
            self.set_font('Helvetica','I',8);self.cell(0,8,s(self.t),0,0,'L');self.cell(0,8,s(ORG),0,1,'R');self.line(10,16,200,16);self.ln(3)
    def footer(self):
        self.set_y(-15);self.set_font('Helvetica','I',8);self.cell(0,10,f'{s(AUTHOR)} | Page {self.page_no()}/{{nb}}',0,0,'C')
    def title_page(self,title,subtitle):
        self.add_page();self.ln(40);self.set_font('Helvetica','B',22);self.cell(0,12,s(title),0,1,'C');self.ln(3)
        self.set_font('Helvetica','',13);self.cell(0,8,s(subtitle),0,1,'C');self.ln(8)
        self.set_font('Helvetica','',11);self.cell(0,7,f'Prepared by {AUTHOR}',0,1,'C');self.cell(0,7,s(ORG),0,1,'C');self.cell(0,7,'March 2026',0,1,'C')
    def h1(self,t): self.ln(4);self.set_font('Helvetica','B',15);self.cell(0,8,s(t),0,1);self.ln(2)
    def h2(self,t): self.ln(3);self.set_font('Helvetica','B',12);self.cell(0,7,s(t),0,1);self.ln(1)
    def h3(self,t): self.ln(2);self.set_font('Helvetica','B',10);self.cell(0,6,s(t),0,1);self.ln(1)
    def p(self,t): self.set_font('Helvetica','',9);self.multi_cell(0,4.5,s(t));self.ln(1.5)
    def bullet(self,t): self.set_font('Helvetica','',9);x=self.get_x();self.cell(8,4.5,'  -',0,0);self.multi_cell(0,4.5,s(t))
    def bold_line(self,t): self.set_font('Helvetica','B',9);self.multi_cell(0,4.5,s(t));self.ln(1)

def make_docx(title, subtitle, sections):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True; run.font.size = Pt(22)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'\nPrepared by {AUTHOR}\n{ORG}\nMarch 2026').font.size = Pt(11)
    
    doc.add_page_break()
    
    for section in sections:
        if section[0] == 'h1':
            doc.add_heading(section[1], level=1)
        elif section[0] == 'h2':
            doc.add_heading(section[1], level=2)
        elif section[0] == 'h3':
            doc.add_heading(section[1], level=3)
        elif section[0] == 'p':
            doc.add_paragraph(section[1])
        elif section[0] == 'bullet':
            doc.add_paragraph(section[1], style='List Bullet')
        elif section[0] == 'bold':
            p = doc.add_paragraph()
            p.add_run(section[1]).bold = True
        elif section[0] == 'break':
            doc.add_page_break()
    
    # Footer
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{ORG} | {AUTHOR}')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128,128,128)
    
    return doc

def make_pdf(title, subtitle, sections):
    pdf = PDF(title)
    pdf.title_page(title, subtitle)
    pdf.add_page()
    
    for section in sections:
        if pdf.get_y() > 270:
            pdf.add_page()
        if section[0] == 'h1': pdf.h1(section[1])
        elif section[0] == 'h2': pdf.h2(section[1])
        elif section[0] == 'h3': pdf.h3(section[1])
        elif section[0] == 'p': pdf.p(section[1])
        elif section[0] == 'bullet': pdf.bullet(section[1])
        elif section[0] == 'bold': pdf.bold_line(section[1])
        elif section[0] == 'break': pdf.add_page()
    return pdf

# ========== ARCHITECTURE DOCUMENT ==========
print("Generating Architecture Document...")
arch_sections = [
    ('h1', '1. Executive Summary'),
    ('p', 'HeatShield Agri is an AI-powered heat stress early warning system for Uganda\'s 12.4 million agricultural workers. The platform delivers real-time WBGT (Wet-Bulb Globe Temperature) predictions and actionable safety guidance through three channels: a web dashboard, an Android app, and a USSD/SMS system for feature phones.'),
    ('p', 'The core scientific engine implements ISO 7243 occupational heat stress calculations consistently across all platforms. A Random Forest machine learning pipeline enhances physics-based predictions with trained ONNX models for temperature, humidity, and wind speed forecasting.'),
    
    ('h1', '2. System Architecture Overview'),
    ('p', 'Multi-channel delivery platform with shared scientific core:'),
    ('bullet', 'Web Dashboard (React + TypeScript) - production deployed at heatshieldagri.vercel.app'),
    ('bullet', 'Android App (Kotlin + Jetpack Compose) - functional prototype with ML inference'),
    ('bullet', 'USSD Backend (Rust + Actix-web) - structured skeleton with full menu system'),
    ('bullet', 'WASM Engine (Rust) - complete calculation library for browser-native performance'),
    ('bullet', 'Random Forest ML Pipeline - ONNX models deployed on web and Android'),
    ('h2', 'Shared Scientific Core'),
    ('bullet', 'ISO 7243 WBGT formula: WBGT = 0.7*Tw + 0.2*Tg + 0.1*Ta'),
    ('bullet', 'Stull (2011) wet-bulb temperature approximation'),
    ('bullet', 'Liljegren globe temperature (Newton-Raphson iterative solver)'),
    ('bullet', 'Risk thresholds: Low (<26C), Moderate (26-28C), High (28-30C), Very High (30-32C), Extreme (>32C)'),
    ('bullet', 'Open-Meteo weather API for real-time data'),
    ('bullet', '12 Uganda districts with GPS coordinates'),
    
    ('h1', '3. Web Dashboard Architecture'),
    ('bold', 'Stack: React 18.2, TypeScript, Vite, TailwindCSS 3.4, Zustand, React Query, Recharts'),
    ('bold', 'Deployment: Vercel (production) | Live URL: https://heatshieldagri.vercel.app'),
    ('h2', 'Key Components'),
    ('bullet', 'Dashboard: Real-time WBGT with current conditions for selected district'),
    ('bullet', 'Forecast: 72-hour WBGT forecast with hourly granularity, BOTH Physics and Random Forest models displayed'),
    ('bullet', 'Heat Map: Interactive SVG map with IDW spatial interpolation across 12 districts'),
    ('bullet', 'Schedule Optimizer: AI-powered work schedule recommendations, daylight hours (6-18)'),
    ('bullet', 'Alerts: SMS alert registration with configurable thresholds, multi-language support'),
    ('bullet', 'Demo: Interactive WBGT calculator and WASM compilation demo'),
    ('bullet', 'Languages: English, Luganda, Runyankole, Acholi, Swahili'),
    ('h2', 'ML Integration'),
    ('bullet', '3 Random Forest ONNX models (~6.2 MB total): temperature, humidity, windspeed'),
    ('bullet', '17-feature engineering pipeline: 8 lag features, 4 cyclical time encodings, 3 rolling stats, 2 delta features'),
    ('bullet', 'Physics-first architecture: NWP forecast -> ISO 7243 baseline -> RF validation (MAE <= 2.0C) -> 70% physics + 30% RF blend'),
    ('bullet', 'ONNX Runtime Web for browser-native inference'),
    
    ('h1', '4. Android App Architecture'),
    ('bold', 'Stack: Kotlin, Jetpack Compose, Material 3, Hilt DI, Coroutines, StateFlow, Retrofit, ONNX Runtime Android'),
    ('bold', 'Min SDK: 26 (Android 8.0) | Target SDK: 34'),
    ('h2', 'Key Screens'),
    ('bullet', 'Dashboard: Real-time weather from Open-Meteo, live WBGT calculation'),
    ('bullet', 'Forecast: Dual-model display showing BOTH Physics and RF WBGT side by side, district selector, daylight hours (6-18)'),
    ('bullet', 'Schedule: ML-powered work schedule optimizer with district selector, work hours selector, productivity scoring'),
    ('bullet', 'Map: Heat risk visualization'),
    ('bullet', 'Demo: Interactive WBGT calculator with sliders'),
    ('h2', 'Architecture Pattern: MVVM'),
    ('bullet', 'ViewModels: ForecastViewModel, ScheduleViewModel (physics-first + RF validation)'),
    ('bullet', 'Repository: WeatherRepository with getWeatherWithHistory() for historical + forecast data'),
    ('bullet', 'ML Engine: HeatShieldMLInference singleton (ONNX Runtime, 2 inference threads)'),
    ('bullet', 'DI: MLModule provides singleton ML inference via Hilt'),
    ('bullet', 'API: WeatherApiService with Retrofit (Open-Meteo, past_days=4 for 96h history)'),
    ('h2', 'ML Integration'),
    ('bullet', '3 ONNX models in assets/models/ (~6.2 MB total)'),
    ('bullet', 'ONNX Runtime Android (native C++ via JNI, high performance)'),
    ('bullet', '17-feature pipeline identical to web'),
    ('bullet', 'Physics-first with RF validation architecture'),
    ('p', 'Branding: "Designed by Wayesu Community Research Organisation Ltd" in Dashboard footer'),
    
    ('h1', '5. USSD Backend Architecture'),
    ('bold', 'Stack: Rust 2021, Actix-web 4, Tokio, SQLx (PostgreSQL), Redis'),
    ('bold', 'Binary: 6.3 MB statically linked | Service Code: *384*HEAT#'),
    ('h2', 'Key Components'),
    ('bullet', 'USSD Callback Handler: Parses Africa\'s Talking format (sessionId, phoneNumber, text)'),
    ('bullet', 'Menu System: 6-item main menu - Heat Risk, Forecast, Safe Hours, SMS Alerts, Location, Language'),
    ('bullet', 'Session Management: State machine with enum states'),
    ('bullet', 'i18n: 5 languages (English, Luganda, Runyankole, Acholi, Swahili), 15+ message keys'),
    ('bullet', 'WBGT Engine: Same ISO 7243 calculations in Rust'),
    ('bullet', 'Location Service: 16 Uganda districts, haversine distance'),
    ('h2', 'Infrastructure (configured, not yet connected)'),
    ('bullet', 'PostgreSQL for user data persistence'),
    ('bullet', 'Redis for session management'),
    ('bullet', 'Africa\'s Talking for SMS delivery'),
    ('bullet', 'Docker multi-stage build (rust:1.74-slim builder, debian:bookworm-slim runtime)'),
    
    ('h1', '6. WASM Engine Architecture'),
    ('bold', 'Stack: Rust, wasm-bindgen, js-sys, web-sys'),
    ('h2', 'Modules'),
    ('bullet', 'wbgt.rs: ISO 7243 WBGT calculation (213 lines)'),
    ('bullet', 'risk.rs: Risk classification with local language messages (250 lines)'),
    ('bullet', 'spatial.rs: IDW spatial interpolation, 5km grid for Uganda (244 lines)'),
    ('bullet', 'schedule.rs: Work schedule optimization (275 lines)'),
    ('bullet', 'Unit tests in each module'),
    ('p', 'Status: Complete source code, JavaScript fallback used in production'),
    
    ('h1', '7. Random Forest ML Pipeline'),
    ('bold', 'Model: Random Forest (scikit-learn trained, exported to ONNX)'),
    ('bold', 'Models: 3 separate models for temperature, humidity, wind speed'),
    ('h2', '17-Feature Engineering Pipeline'),
    ('bullet', 'Lag features: t-1, t-2, t-3, t-6, t-12, t-24, t-48, t-72 hours'),
    ('bullet', 'Cyclical encodings: sin/cos of hour-of-day, sin/cos of day-of-year'),
    ('bullet', 'Rolling statistics: mean_24h, mean_72h, std_24h'),
    ('bullet', 'Delta features: delta_1h, delta_24h'),
    ('h2', 'Deployment'),
    ('bullet', 'Web: ONNX Runtime Web (JavaScript)'),
    ('bullet', 'Android: ONNX Runtime Android (native C++ via JNI)'),
    ('bullet', 'Architecture: Physics-first - NWP forecast provides baseline; RF validates when MAE <= 2.0C'),
    
    ('h1', '8. Cross-Platform Consistency'),
    ('p', 'All platforms implement identical scientific calculations:'),
    ('bullet', 'WBGT Formula: 0.7*Tw + 0.2*Tg + 0.1*Ta (all 4 platforms)'),
    ('bullet', 'Wet-bulb method: Stull 2011 (all 4 platforms)'),
    ('bullet', 'Globe temp method: Liljegren Newton-Raphson (all 4 platforms)'),
    ('bullet', 'Risk thresholds: Low 26C, Moderate 28C, High 30C, Very High 32C (all 4 platforms)'),
    ('bullet', 'Weather API: Open-Meteo (Web + Android live, USSD not yet connected)'),
    ('p', 'This consistency ensures all platforms give the same risk assessment for the same conditions.'),
    
    ('h1', '9. Data Flow'),
    ('bold', 'Weather Data Source: Open-Meteo API (ECMWF + GFS models)'),
    ('h2', 'Processing Pipeline'),
    ('bullet', '1. Fetch weather data from Open-Meteo (current + 96h history + 48h forecast)'),
    ('bullet', '2. Compute physics WBGT using ISO 7243'),
    ('bullet', '3. If RF models available: run 17-feature pipeline, predict temp/hum/wind, compute RF WBGT'),
    ('bullet', '4. Compare RF vs Physics (calculate MAE)'),
    ('bullet', '5. Display both models side by side for comparison'),
    
    ('h1', '10. Security and Deployment'),
    ('bullet', 'Web: Vercel with security headers, SPA routing, asset caching, HTTPS'),
    ('bullet', 'Android: Hilt DI, Proguard minification, release signing'),
    ('bullet', 'USSD: Docker containerization, non-root user, TLS, env var secrets'),
    ('bullet', 'Data: Minimal collection (phone number + district only), TLS encryption, no data sales'),
]

arch_pdf = make_pdf('HeatShield Agri', 'Platform Architecture Document', arch_sections)
arch_pdf.output(os.path.join(DOCS_DIR, 'HeatShield_Agri_Architecture.pdf'))
print("  Architecture PDF generated")

arch_docx = make_docx('HeatShield Agri', 'Platform Architecture Document', arch_sections)
arch_docx.save(os.path.join(DOCS_DIR, 'HeatShield_Agri_Architecture.docx'))
print("  Architecture DOCX generated")

# ========== STRATEGY REPORT PDF + DOCX ==========
print("Generating Strategy Report PDF + DOCX...")
md_path = os.path.join(DOCS_DIR, 'Deployment_Readiness_Strategy_Report.md')
with open(md_path, 'r') as f:
    md_content = f.read()

# Parse MD into sections
strat_sections = []
for line in md_content.split('\n'):
    line = line.strip()
    if not line or line == '---': continue
    if line.startswith('### '): strat_sections.append(('h2', line[4:].strip('*').strip()))
    elif line.startswith('## '): strat_sections.append(('h1', line[3:].strip('*').strip()))
    elif line.startswith('# '): strat_sections.append(('h1', line[2:].strip('*').strip()))
    elif line.startswith('**') and line.endswith('**'): strat_sections.append(('bold', line.strip('*').strip()))
    elif line.startswith('- ') or line.startswith('* '): strat_sections.append(('bullet', line[2:].strip()))
    elif line.startswith('| '): strat_sections.append(('p', line))  # table rows as text
    else: strat_sections.append(('p', line))

strat_pdf = make_pdf('HeatShield Agri', 'Deployment Readiness Strategy Report', strat_sections)
strat_pdf.output(os.path.join(DOCS_DIR, 'Deployment_Readiness_Strategy_Report.pdf'))
print("  Strategy Report PDF generated")

strat_docx = make_docx('HeatShield Agri', 'Deployment Readiness Strategy Report', strat_sections)
strat_docx.save(os.path.join(DOCS_DIR, 'Deployment_Readiness_Strategy_Report.docx'))
print("  Strategy Report DOCX generated")

# ========== QA PROFESSIONAL ==========
print("Generating QA Professional PDF + DOCX...")

# Read existing docx and replace XGBoost
try:
    qa_doc = Document(os.path.join(DOCS_DIR, 'HeatShield_QA_Professional.docx'))
    for para in qa_doc.paragraphs:
        for run in para.runs:
            if 'XGBoost' in run.text:
                run.text = run.text.replace('XGBoost', 'Random Forest')
            if 'xgboost' in run.text.lower():
                run.text = run.text.replace('xgboost', 'Random Forest').replace('Xgboost', 'Random Forest')
            if 'Wayesu Research Ltd' in run.text:
                run.text = run.text.replace('Wayesu Research Ltd', ORG)
    
    # Add org info to header section if possible
    for section in qa_doc.sections:
        footer = section.footer
        if footer.paragraphs:
            for p in footer.paragraphs:
                for run in p.runs:
                    if 'XGBoost' in run.text:
                        run.text = run.text.replace('XGBoost', 'Random Forest')
    
    qa_doc.save(os.path.join(DOCS_DIR, 'HeatShield_QA_Professional.docx'))
    print("  QA Professional DOCX updated (XGBoost -> Random Forest)")
except Exception as e:
    print(f"  Warning: Could not update QA Professional DOCX: {e}")

# Generate QA Professional PDF from scratch
qa_sections = [
    ('h1', 'HEATSHIELD AGRI - Interview Q&A Reference'),
    ('bold', 'AI-Powered Heat Stress Early Warning System'),
    ('bold', f'Prepared for: {AUTHOR}, {ORG}'),
    ('bold', 'Interview: March 9, 2026, 21:00 EAT'),
    ('p', 'Format: 5 min intro + 10 min pitch + 30 min Q&A'),
    ('break', ''),
    
    ('h1', '1. Technical & AI Questions'),
    ('h2', 'How accurate are your WBGT predictions?'),
    ('p', 'Our predictions use the ISO 7243 WBGT calculation standard - the internationally recognized method for occupational heat assessment. We combine multiple data sources including Open-Meteo weather forecasts (ECMWF and GFS models), SRTM terrain data, and spatial interpolation algorithms. We display uncertainty bounds with every prediction and use conservative thresholds that err toward caution. During the pilot phase, we\'ll collect ground-truth data to quantify actual prediction error.'),
    
    ('h2', 'Why did you choose Random Forest with a physics-first architecture?'),
    ('p', 'This hybrid approach captures the best of both worlds. Random Forest excels at handling tabular weather data with complex feature interactions (temperature, humidity, wind, terrain), while the physics layer ensures predictions respect known thermodynamic relationships from ISO 7243. The physics-first constraint prevents the model from making physically impossible predictions - a common failure mode of pure ML approaches.'),
    ('p', 'Random Forest is computationally efficient and deployable via ONNX Runtime without expensive GPU infrastructure, which is essential for sustainable operations. On Android, ONNX Runtime uses native C++ via JNI for high-performance inference. On the web, ONNX Runtime Web runs directly in the browser. We have deployed 3 trained models (temperature, humidity, wind speed) with a 17-feature engineering pipeline including lag features, cyclical time encodings, rolling statistics, and delta features.'),
    
    ('h2', 'How does the work schedule optimizer function?'),
    ('p', 'The optimizer analyzes the WBGT forecast and identifies all hours below the recommended work threshold for moderate agricultural work (ISO 7243 categories). It filters to daylight hours (6:00-18:00) and applies practical constraints: farmers need 6-8 productive hours and prefer contiguous blocks. The output is specific recommendations like "6:00 AM - 11:00 AM, then 4:00 PM - 6:00 PM" with estimated safe work hours and productivity scores.'),
    
    ('h2', 'What happens when predictions are wrong?'),
    ('p', 'We design for this scenario explicitly. Every alert includes guidance to "listen to your body" and recognize heat illness symptoms. We use conservative thresholds that minimize false negatives. We implement a feedback loop where users report via SMS whether conditions matched predictions. The pilot phase explicitly validates prediction accuracy before scaling.'),
    ('break', ''),
    
    ('h1', '2. Impact & Reach Questions'),
    ('h2', 'How will you reach farmers without smartphones?'),
    ('p', 'USSD/SMS is a primary channel, not an afterthought. USSD works on any basic feature phone without internet. Farmers dial *384*HEAT# and navigate simple numbered menus. SMS alerts are in local languages: "HIGH HEAT TODAY. Work 5AM-10AM only. Rest midday. Drink water every 30 min." For farmers without any phone, extension workers share information through community gatherings.'),
    
    ('h2', 'How will you reach 10,000 farmers in 12 months?'),
    ('p', 'We plug into existing infrastructure with trusted relationships. UNFFE\'s cooperative networks become distribution partners. We train 10 extension workers who each work with ~100 farmer groups. Once registered, SMS delivery is automated. Growth: Months 1-3 = 500, Months 4-6 = 2,000, Months 7-12 = 10,000.'),
    
    ('h2', 'How do you ensure equitable reach for women farmers?'),
    ('p', 'Women face higher heat stress risk (PSI 3.28 vs 2.9 for men) but often have less phone access. We partner with women\'s farming groups. For shared phones, cooperative-level alerts ensure women learn through group meetings. We track engagement by gender and adjust outreach.'),
    ('break', ''),
    
    ('h1', '3. Team & Execution Questions'),
    ('h2', 'Can you execute this as a solo founder?'),
    ('p', 'The MVP was built solo, demonstrating execution capability. But the 12-month plan doesn\'t rely on me alone. I handle technical development, ML models, architecture, partnerships, and fundraising. The $47K personnel budget funds a part-time developer and extension workers. Partners contribute: UNFFE provides cooperative access, extension workers do farmer training, Makerere provides research support.'),
    
    ('h2', 'What qualifies you to lead this?'),
    ('p', f'{AUTHOR} brings: 10+ years as biostatistician at UVRI; MSc Statistics from Linkoping University; current PhD in ML/Computer Vision; technical proficiency in R, Python, Rust, TypeScript. Built the entire HeatShield MVP. Unique intersection of biostatistics + ML + Uganda health systems expertise.'),
    ('break', ''),
    
    ('h1', '4. Sustainability & Scale Questions'),
    ('h2', 'What\'s your revenue model after the grant?'),
    ('p', 'Basic alerts to individual farmers remain FREE forever. B2B dashboards ($50-200/month) for cooperatives and export companies. Institutional partnerships with MAAIF, World Bank, USAID. Climate adaptation funding from GCF, Adaptation Fund. Cost: $0.30/farmer/year for SMS.'),
    
    ('h2', 'How does this scale beyond Uganda?'),
    ('p', 'Architecture designed for replication from day one. Weather data sources cover all East Africa. WBGT calculations are universal physics. East Africa has 100M+ agricultural workers. Open-source (MIT license) enables others to deploy.'),
    ('break', ''),
    
    ('h1', '5. Responsible AI Questions'),
    ('h2', 'How do you handle farmer data privacy?'),
    ('p', 'Data minimization: only phone number and district. No health data collected. Phone numbers encrypted. Users opt-in with local language explanation. Easy opt-out via SMS. Never sold to third parties.'),
    
    ('h2', 'What about environmental impact of your AI?'),
    ('p', 'Lightweight by design: Random Forest and spatial interpolation, not GPU-intensive deep learning. Pre-calculated forecasts are cached and served. Environmental cost is negligible compared to climate adaptation benefits. Cloud providers have renewable energy commitments.'),
    ('break', ''),
    
    ('h1', 'Key Messages Summary'),
    ('bullet', 'Working MVP: Live at heatshieldagri.vercel.app - proof of execution, not just plans'),
    ('bullet', 'Unserved Market: 12.4M agricultural workers with zero existing heat protection'),
    ('bullet', 'Proven Science: ISO 7243 standard + Random Forest ML - innovation is application'),
    ('bullet', 'Inclusive Design: USSD/SMS for feature phones, local languages, gender equity focus'),
    ('bullet', 'Regional Scale: Uganda proves concept for 100M+ East African workers'),
    ('bullet', 'Sustainability: Freemium model - free farmer alerts, paid institutional dashboards'),
    ('bullet', f'Founder: {AUTHOR}, {ORG}'),
]

qa_pdf = make_pdf('HeatShield Agri', 'Interview Q&A Reference', qa_sections)
qa_pdf.output(os.path.join(DOCS_DIR, 'HeatShield_QA_Professional.pdf'))
print("  QA Professional PDF generated")

# ========== QA TALKING POINTS ==========
print("Generating QA Talking Points PDF + DOCX...")

try:
    tp_doc = Document(os.path.join(DOCS_DIR, 'HeatShield_QA_Talking_Points.docx'))
    for para in tp_doc.paragraphs:
        for run in para.runs:
            if 'XGBoost' in run.text:
                run.text = run.text.replace('XGBoost', 'Random Forest')
            if 'Wayesu Research Ltd' in run.text:
                run.text = run.text.replace('Wayesu Research Ltd', ORG)
    for table in tp_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if 'XGBoost' in run.text:
                            run.text = run.text.replace('XGBoost', 'Random Forest')
    tp_doc.save(os.path.join(DOCS_DIR, 'HeatShield_QA_Talking_Points.docx'))
    print("  QA Talking Points DOCX updated")
except Exception as e:
    print(f"  Warning: Could not update QA Talking Points DOCX: {e}")

# Generate Talking Points PDF
tp_sections = [
    ('h1', 'ACTIVATE AI CHALLENGE - Finalist Interview Preparation'),
    ('bold', f'Detailed Talking Points for {AUTHOR}'),
    ('bold', f'{ORG}'),
    ('bold', 'Interview: March 9, 2026, 21:00 EAT'),
    ('p', 'Structure: 5 min intro + 10 min pitch + 30 min Q&A'),
    ('break', ''),
    
    ('h1', 'SECTION 1: TECHNICAL & AI QUESTIONS'),
    ('h2', 'Q: How accurate are your WBGT predictions?'),
    ('bold', 'Talking Points:'),
    ('bullet', '1. We use the ISO 7243 WBGT calculation standard - internationally recognized for occupational heat assessment.'),
    ('bullet', '2. Multiple data sources: Open-Meteo (ECMWF + GFS models), SRTM terrain data, spatial interpolation (IDW).'),
    ('bullet', '3. Current MVP validation: Cross-reference against expected values from known conditions. Physics-based approach means error comes primarily from upstream weather forecast uncertainty.'),
    ('bullet', '4. Planned improvements: ICPAC regional forecast integration, ground-truth data collection during pilot.'),
    ('bullet', '5. We display uncertainty bounds and use conservative thresholds that err toward caution.'),
    ('bold', 'Key Message: Predictions based on international standards and multiple data sources. Transparent about uncertainty.'),
    
    ('h2', 'Q: Why did you choose Random Forest with a physics-first architecture?'),
    ('bold', 'Talking Points:'),
    ('bullet', '1. The hybrid approach captures the best of both worlds: Random Forest excels at tabular data with complex feature interactions, while physics constraints ensure ISO 7243 compliance.'),
    ('bullet', '2. We have DEPLOYED 3 trained Random Forest models (temperature, humidity, wind speed) as ONNX files (~6.2 MB total). They run on Android via ONNX Runtime (native C++ via JNI) and on web via ONNX Runtime Web.'),
    ('bullet', '3. Our 17-feature engineering pipeline uses 8 lag features (1-72h), 4 cyclical time encodings, 3 rolling statistics, and 2 delta features for robust predictions.'),
    ('bullet', '4. Random Forest is computationally efficient - no GPU needed. Deployable on standard infrastructure including mobile devices.'),
    ('bullet', '5. This architecture is well-established in climate science. We apply proven methods to an underserved use case.'),
    ('bold', 'Key Message: We chose proven, efficient methods. Random Forest gives interpretability for stakeholders, ONNX gives cross-platform native performance.'),
    
    ('h2', 'Q: How does the AI work schedule optimizer function?'),
    ('bold', 'Talking Points:'),
    ('bullet', '1. Analyzes WBGT forecast, identifies hours below recommended threshold for moderate agricultural work (ISO 7243).'),
    ('bullet', '2. Filters to daylight hours (6:00-18:00), applies farmer constraints (6-8 hours needed, contiguous blocks).'),
    ('bullet', '3. Outputs specific windows: e.g., "6:00 AM - 11:00 AM, then 4:00 PM - 6:00 PM" with productivity scores.'),
    ('bullet', '4. This is a constraint satisfaction problem processing thousands of forecast data points humans can\'t manually analyze.'),
    ('bold', 'Key Message: Practical and explainable. AI in the sense of automated intelligent analysis, not black-box magic.'),
    
    ('h2', 'Q: Is the AI/ML model trained and deployed?'),
    ('bold', 'Talking Points:'),
    ('bullet', '1. YES - we have trained and deployed Random Forest ML models on BOTH the web dashboard and Android app.'),
    ('bullet', '2. Three ONNX models predict temperature, humidity, and wind speed using a 17-feature engineering pipeline.'),
    ('bullet', '3. The system uses physics-first architecture: ISO 7243 provides baseline WBGT, Random Forest validates and enhances when deviation is within 2.0 degrees C.'),
    ('bullet', '4. Both the Forecast and Schedule tabs on Android show both Physics and RF model results side by side.'),
    ('bullet', '5. The web dashboard similarly displays both models for comparison and transparency.'),
    ('bold', 'Key Message: ML is not planned - it is DEPLOYED and running. You can see both models compared on the live website and Android app.'),
    ('break', ''),
    
    ('h1', 'SECTION 2: IMPACT & REACH QUESTIONS'),
    ('h2', 'Q: How will you reach farmers who don\'t have smartphones?'),
    ('bold', 'Talking Points:'),
    ('bullet', '1. USSD/SMS is a primary channel. Works on any $15 feature phone without internet.'),
    ('bullet', '2. Farmers dial *384*HEAT# for simple menu. SMS alerts in local languages.'),
    ('bullet', '3. Community amplification through extension workers and cooperative leaders.'),
    ('bullet', '4. No internet required - USSD/SMS works on 2G networks covering nearly all of Uganda.'),
    ('bold', 'Key Message: We meet farmers where they are - feature phones, no internet, local languages.'),
    
    ('h2', 'Q: How do you measure actual health impact?'),
    ('bold', 'Talking Points:'),
    ('bullet', '1. Behavior change is primary outcome - measurable and causally linked to health protection.'),
    ('bullet', '2. Self-reported work schedule changes via monthly SMS check-ins and surveys.'),
    ('bullet', '3. Baseline/endline surveys measuring heat illness symptoms. Compare users vs non-users.'),
    ('bullet', '4. Makerere University partnership ensures rigorous evaluation methodology.'),
    ('bold', 'Key Message: Measure the full chain: alert delivery -> behavior change -> symptom reduction.'),
    ('break', ''),
    
    ('h1', 'SECTION 3: TEAM & EXECUTION'),
    ('h2', 'Q: Can you execute this as a solo founder?'),
    ('bold', 'Talking Points:'),
    ('bullet', f'1. The MVP was built solo by {AUTHOR}, demonstrating execution capability.'),
    ('bullet', '2. 12-month plan leverages partners: UNFFE (cooperative access), extension workers (training), Makerere (research).'),
    ('bullet', '3. $47K personnel budget funds part-time developer and extension worker payments.'),
    ('bullet', '4. System scales through automation (SMS) and existing human infrastructure.'),
    ('bold', 'Key Message: Solo founder does not equal solo execution. The MVP proves capability; the grant funds a team.'),
    
    ('h2', 'Q: What\'s your background?'),
    ('bold', 'Talking Points:'),
    ('bullet', f'1. {AUTHOR}: 10+ years biostatistician at Uganda Virus Research Institute (UVRI).'),
    ('bullet', '2. MSc Statistics from Linkoping University (Sweden).'),
    ('bullet', '3. Current PhD in Machine Learning/Computer Vision.'),
    ('bullet', '4. Technical: R, Python, Rust, TypeScript. Built entire HeatShield MVP.'),
    ('bullet', '5. Unique intersection: health systems + statistics + ML + Uganda context.'),
    ('bold', f'Key Message: {AUTHOR} combines statistical rigor, ML capability, health research, and Uganda context.'),
    ('break', ''),
    
    ('h1', 'SECTION 4: SUSTAINABILITY & SCALE'),
    ('h2', 'Q: Revenue model after the grant?'),
    ('bullet', 'Basic farmer alerts: FREE forever (public good)'),
    ('bullet', 'B2B dashboards: $50-200/month for cooperatives, export companies'),
    ('bullet', 'Institutional partners: MAAIF, World Bank, USAID contracts'),
    ('bullet', 'Climate finance: GCF, Adaptation Fund grants'),
    ('bullet', 'Cost advantage: $0.30/farmer/year for SMS'),
    
    ('h2', 'Q: How does this scale beyond Uganda?'),
    ('bullet', 'Weather data (Open-Meteo, ICPAC) covers all East Africa'),
    ('bullet', 'WBGT calculations are universal physics'),
    ('bullet', '100M+ East African agricultural workers face similar challenges'),
    ('bullet', 'Open source (MIT license) enables replication'),
    ('break', ''),
    
    ('h1', 'SECTION 5: RESPONSIBLE AI'),
    ('h2', 'Q: How do you handle data privacy?'),
    ('bullet', 'Data minimization: only phone number and district'),
    ('bullet', 'Phone numbers encrypted, TLS for all transmission'),
    ('bullet', 'Active opt-in with local language explanation'),
    ('bullet', 'Never sold to third parties'),
    
    ('h2', 'Q: Environmental impact of your AI?'),
    ('bullet', 'Lightweight: Random Forest and spatial interpolation, not GPU-intensive deep learning'),
    ('bullet', 'Pre-calculated forecasts cached and served'),
    ('bullet', 'Environmental cost negligible vs climate adaptation benefits'),
    ('break', ''),
    
    ('h1', 'KEY THEMES TO EMPHASIZE'),
    ('bullet', 'Working MVP: Live at heatshieldagri.vercel.app - proof of execution'),
    ('bullet', 'Random Forest ML: DEPLOYED on web and Android, not just planned'),
    ('bullet', 'Unserved population: 12.4M workers with zero heat protection'),
    ('bullet', 'Proven science: ISO 7243 + Random Forest - innovation is application'),
    ('bullet', 'Inclusive: USSD/SMS for feature phones, 5 local languages'),
    ('bullet', 'Regional scalability: Uganda is proof of concept for 100M+ workers'),
    ('p', ''),
    ('bold', f'{AUTHOR} | {ORG}'),
    ('bold', 'Remember: You\'ve built something real. Show them.'),
]

tp_pdf = make_pdf('HeatShield Agri', 'Interview Preparation - Detailed Talking Points', tp_sections)
tp_pdf.output(os.path.join(DOCS_DIR, 'HeatShield_QA_Talking_Points.pdf'))
print("  QA Talking Points PDF generated")

print("\n=== ALL DOCUMENTS GENERATED ===")
for f in ['HeatShield_Agri_Architecture.pdf', 'HeatShield_Agri_Architecture.docx',
          'Deployment_Readiness_Strategy_Report.pdf', 'Deployment_Readiness_Strategy_Report.docx',
          'HeatShield_QA_Professional.pdf', 'HeatShield_QA_Professional.docx',
          'HeatShield_QA_Talking_Points.pdf', 'HeatShield_QA_Talking_Points.docx']:
    path = os.path.join(DOCS_DIR, f)
    if os.path.exists(path):
        size = os.path.getsize(path)
        print(f"  OK {f} ({size:,} bytes)")
    else:
        print(f"  MISSING {f}")
